from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Cm, Pt, RGBColor


def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), "true")
    tr_pr.append(tbl_header)
    return row


def insert_hr(paragraph):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    p_pr.insert_element_before(
        p_bdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    p_bdr.append(bottom)


def fecha_numerica_a_fecha(fecha):
    array_fecha = fecha.split("/")

    if int(array_fecha[0]) != 0:
        match int(array_fecha[1]):
            case 1:
                return str(str(int(array_fecha[0])) + " enero " + str(int(array_fecha[2])))
            case 2:
                return str(str(int(array_fecha[0])) + " febrero " + str(int(array_fecha[2])))
            case 3:
                return str(str(int(array_fecha[0])) + " marzo " + str(int(array_fecha[2])))
            case 4:
                return str(str(int(array_fecha[0])) + " abril " + str(int(array_fecha[2])))
            case 5:
                return str(str(int(array_fecha[0])) + " mayo " + str(int(array_fecha[2])))
            case 6:
                return str(str(int(array_fecha[0])) + " junio " + str(int(array_fecha[2])))
            case 7:
                return str(str(int(array_fecha[0])) + " julio " + str(int(array_fecha[2])))
            case 8:
                return str(str(int(array_fecha[0])) + " agosto " + str(int(array_fecha[2])))
            case 9:
                return str(str(int(array_fecha[0])) + " septiembre " + str(int(array_fecha[2])))
            case 10:
                return str(str(int(array_fecha[0])) + " octubre " + str(int(array_fecha[2])))
            case 11:
                return str(str(int(array_fecha[0])) + " noviembre " + str(int(array_fecha[2])))
            case 12:
                return str(str(int(array_fecha[0])) + " diciembre " + str(int(array_fecha[2])))
    else:
        match int(array_fecha[1]):
            case 1:
                return str("Enero de " + array_fecha[2])
            case 2:
                return str("Febrero de " + array_fecha[2])
            case 3:
                return str("Marzo de " + array_fecha[2])
            case 4:
                return str("Abril de " + array_fecha[2])
            case 5:
                return str("Mayo de " + array_fecha[2])
            case 6:
                return str("Junio de " + array_fecha[2])
            case 7:
                return str("Julio de " + array_fecha[2])
            case 8:
                return str("Agosto de " + array_fecha[2])
            case 9:
                return str("Septiembre de " + array_fecha[2])
            case 10:
                return str("Octubre de " + array_fecha[2])
            case 11:
                return str("Noviembre de " + array_fecha[2])
            case 12:
                return str("Diciembre de " + array_fecha[2])


def comprovar_plazo(fecha_pedido, fecha_plazo):
    array_fecha = fecha_plazo.split("/")
    array_fecha_pedido = fecha_pedido.split("/")
    print(array_fecha)
    print(array_fecha_pedido)

    if array_fecha[0] != '00' and array_fecha[1] != '00' and array_fecha[2] == '0000':
        return str(int(array_fecha[0])) + ' / ' + str(int(array_fecha[1])) + ' dias'

    elif array_fecha[0] == '00' and array_fecha[1] == '00' and array_fecha[2] == '0000':
        return "A CONVENIR"

    elif array_fecha[1] == '00' and array_fecha[2] == '0000':
        if int(array_fecha[0]) == 1:
            return str(int(array_fecha[0])) + ' dia/day'
        else:
            return str(int(array_fecha[0])) + ' dias/days'

    elif array_fecha[0] == '00' and array_fecha[2] == '0000':
        if int(array_fecha[1]) == 1:
            return str(int(array_fecha[1])) + ' mes/month'
        else:
            return str(int(array_fecha[1])) + ' meses/months'

    elif array_fecha[1] != '00' and array_fecha[2] != '0000':
        if int(array_fecha[1]) > int(array_fecha_pedido[1]) and int(array_fecha[2]) == int(array_fecha_pedido[2]) or int(array_fecha[2]) > int(array_fecha_pedido[2]):
            return fecha_numerica_a_fecha(fecha_plazo)
        else:
            return "[STOCK]"


def comprovar_stock(fecha_pedido, fecha_plazo):

    if fecha_pedido >= fecha_plazo or str(fecha_plazo) == "00/00/0000":
        return True
    else:
        array_fecha_pedido = str(fecha_pedido).split('/')
        array_fecha_plazo = str(fecha_plazo).split('/')

        if int(str(array_fecha_pedido[0]).strip()) == int(str(array_fecha_plazo[0]).strip()) and \
                int(str(array_fecha_pedido[1]).strip()) == int(str(array_fecha_plazo[1]).strip()) and \
                int(str(array_fecha_pedido[2]).strip()) == int(str(array_fecha_plazo[2]).strip()):
            return True
        else:
            return False


def comprovar_usuario(usuario):
    match usuario:
        case 'lai':
            return 'Laia'
        case 'jor':
            return 'Jordi'
        case 'car':
            return 'Carmen'
        case 'jos':
            return 'José'
        case 'ant':
            return 'Antonio'
        case 'lou':
            return 'Lourdes'
        case 'ang':
            return 'Angeles'
        case 'luz':
            return 'Mari Luz'
        case 'jua':
            return 'Juan Manuel'
        case _:
            return 'Lusan'


def crear_tabla_clientes(doc):
    table = doc.add_table(rows=1, cols=6)

    table.allow_autofit = False

    for i in range(6):
        for cell in table.columns[i].cells:
            if i == 0:
                cell.width = Cm(7)
            elif i == 1:
                cell.width = Cm(14.75)
            elif i == 2:
                cell.width = Cm(2.25)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif i == 3:
                cell.width = Cm(2)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif i == 4:
                cell.width = Cm(1.25)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif i == 5:
                cell.width = Cm(2)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    hdr = table.rows[0]
    hdr_cells = hdr.cells

    hdr_cells[0].paragraphs[0].add_run('REF.\n').font.size = Pt(9)
    hdr_cells[0].paragraphs[0].add_run('REF.\n').font.size = Pt(9)
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[0].paragraphs[0].runs[1].font.italic = True

    hdr_cells[1].paragraphs[0].add_run('DESCRIPCION\n').font.size = Pt(9)
    hdr_cells[1].paragraphs[0].add_run('SPECIFICATION\n').font.size = Pt(9)
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].paragraphs[0].runs[1].font.italic = True
    hdr_cells[1].paragraphs[0].runs[1].font.bold = False

    hdr_cells[2].paragraphs[0].add_run('CANTIDAD\n').font.size = Pt(9)
    hdr_cells[2].paragraphs[0].add_run('QUANTITY\n').font.size = Pt(9)
    hdr_cells[2].paragraphs[0].runs[0].font.bold = True
    hdr_cells[2].paragraphs[0].runs[1].font.italic = True
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    hdr_cells[3].paragraphs[0].add_run('PRECIO\n').font.size = Pt(9)
    hdr_cells[3].paragraphs[0].add_run('PRICE\n').font.size = Pt(9)
    hdr_cells[3].paragraphs[0].add_run('EUROx100').font.size = Pt(9)
    hdr_cells[3].paragraphs[0].runs[0].font.bold = True
    hdr_cells[3].paragraphs[0].runs[1].font.italic = True
    hdr_cells[3].paragraphs[0].runs[2].font.bold = True
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    hdr_cells[4].paragraphs[0].add_run('DTO.\n').font.size = Pt(9)
    hdr_cells[4].paragraphs[0].add_run('DIS.\n').font.size = Pt(9)
    hdr_cells[4].paragraphs[0].add_run('%').font.size = Pt(9)
    hdr_cells[4].paragraphs[0].runs[0].font.bold = True
    hdr_cells[4].paragraphs[0].runs[1].font.italic = True
    hdr_cells[4].paragraphs[0].runs[2].font.bold = True
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    hdr_cells[5].paragraphs[0].add_run('IMPORTE\n').font.size = Pt(9)
    hdr_cells[5].paragraphs[0].add_run('AMOUNT\n').font.size = Pt(9)
    hdr_cells[5].paragraphs[0].add_run('EURO').font.size = Pt(9)
    hdr_cells[5].paragraphs[0].runs[0].font.bold = True
    hdr_cells[5].paragraphs[0].runs[1].font.italic = True
    hdr_cells[5].paragraphs[0].runs[2].font.bold = True
    hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    hdr.height = Cm(1.25)
    hdr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    barra_cabeza = table.add_row()
    barra_cabeza_tabla = barra_cabeza.cells

    barra_cabeza_tabla[5].merge(barra_cabeza_tabla[4])
    barra_cabeza_tabla[4].merge(barra_cabeza_tabla[3])
    barra_cabeza_tabla[3].merge(barra_cabeza_tabla[2])
    barra_cabeza_tabla[2].merge(barra_cabeza_tabla[1])
    barra_cabeza_tabla[1].merge(barra_cabeza_tabla[0])

    barra_cabeza.height = Cm(0.65)
    barra_cabeza.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    insert_hr(barra_cabeza_tabla[0].paragraphs[0])

    set_repeat_table_header(table.rows[0])
    set_repeat_table_header(table.rows[1])

    return table


def crear_tabla_consulta(doc):
    table = doc.add_table(rows=1, cols=4)

    for i in range(4):
        for cell in table.columns[i].cells:
            if i == 0:
                cell.width = Inches(4)
            elif i == 1:
                cell.width = Inches(1)
            elif i == 2:
                cell.width = Inches(0.5)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif i == 3:
                cell.width = Inches(0.5)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    hdr = table.rows[0]
    hdr_cells = hdr.cells

    hdr_cells[0].paragraphs[0].add_run('DESCRIPCION\n').font.size = Pt(9)
    hdr_cells[0].paragraphs[0].add_run('SPECIFICATION\n').font.size = Pt(9)
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[0].paragraphs[0].runs[1].font.italic = True
    hdr_cells[0].paragraphs[0].runs[1].font.bold = False

    hdr_cells[1].paragraphs[0].add_run('CANTIDAD\n').font.size = Pt(9)
    hdr_cells[1].paragraphs[0].add_run('QUANTITY\n').font.size = Pt(9)
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].paragraphs[0].runs[1].font.italic = True
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    hdr_cells[2].paragraphs[0].add_run('PRECIO\n').font.size = Pt(9)
    hdr_cells[2].paragraphs[0].add_run('PRICE\n').font.size = Pt(9)
    hdr_cells[2].paragraphs[0].add_run('EUROx100').font.size = Pt(9)
    hdr_cells[2].paragraphs[0].runs[0].font.bold = True
    hdr_cells[2].paragraphs[0].runs[1].font.italic = True
    hdr_cells[2].paragraphs[0].runs[2].font.bold = True
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    hdr_cells[3].paragraphs[0].add_run('DTO.\n').font.size = Pt(9)
    hdr_cells[3].paragraphs[0].add_run('DIS.\n').font.size = Pt(9)
    hdr_cells[3].paragraphs[0].add_run('%').font.size = Pt(9)
    hdr_cells[3].paragraphs[0].runs[0].font.bold = True
    hdr_cells[3].paragraphs[0].runs[1].font.italic = True
    hdr_cells[3].paragraphs[0].runs[2].font.bold = True
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    hdr.height = Cm(1.25)
    hdr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    barra_cabeza = table.add_row()
    barra_cabeza_tabla = barra_cabeza.cells

    barra_cabeza_tabla[3].merge(barra_cabeza_tabla[2])
    barra_cabeza_tabla[2].merge(barra_cabeza_tabla[1])
    barra_cabeza_tabla[1].merge(barra_cabeza_tabla[0])

    barra_cabeza.height = Cm(0.65)
    barra_cabeza.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    insert_hr(barra_cabeza_tabla[0].paragraphs[0])

    set_repeat_table_header(table.rows[0])
    set_repeat_table_header(table.rows[1])

    return table


def crear_tabla_pedprov(doc):
    table = doc.add_table(rows=1, cols=5)

    for i in range(5):
        for cell in table.columns[i].cells:
            if i == 0:
                cell.width = Inches(4)
            elif i == 1:
                cell.width = Inches(1)
            elif i == 2:
                cell.width = Inches(0.5)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif i == 3:
                cell.width = Inches(0.5)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif i == 4:
                cell.width = Inches(1.5)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    hdr = table.rows[0]
    hdr_cells = hdr.cells

    hdr_cells[0].paragraphs[0].add_run('DESCRIPCION\n').font.size = Pt(9)
    hdr_cells[0].paragraphs[0].add_run('SPECIFICATION\n').font.size = Pt(9)
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[0].paragraphs[0].runs[1].font.italic = True
    hdr_cells[0].paragraphs[0].runs[1].font.bold = False

    hdr_cells[1].paragraphs[0].add_run('CANTIDAD\n').font.size = Pt(9)
    hdr_cells[1].paragraphs[0].add_run('QUANTITY\n').font.size = Pt(9)
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].paragraphs[0].runs[1].font.italic = True
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    hdr_cells[2].paragraphs[0].add_run('PRECIO\n').font.size = Pt(9)
    hdr_cells[2].paragraphs[0].add_run('PRICE\n').font.size = Pt(9)
    hdr_cells[2].paragraphs[0].add_run('EUROx100').font.size = Pt(9)
    hdr_cells[2].paragraphs[0].runs[0].font.bold = True
    hdr_cells[2].paragraphs[0].runs[1].font.italic = True
    hdr_cells[2].paragraphs[0].runs[2].font.bold = True
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    hdr_cells[3].paragraphs[0].add_run('DTO.\n').font.size = Pt(9)
    hdr_cells[3].paragraphs[0].add_run('DIS.\n').font.size = Pt(9)
    hdr_cells[3].paragraphs[0].add_run('%').font.size = Pt(9)
    hdr_cells[3].paragraphs[0].runs[0].font.bold = True
    hdr_cells[3].paragraphs[0].runs[1].font.italic = True
    hdr_cells[3].paragraphs[0].runs[2].font.bold = True
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    hdr_cells[4].paragraphs[0].add_run('IMPORTE\n').font.size = Pt(9)
    hdr_cells[4].paragraphs[0].add_run('AMOUNT\n').font.size = Pt(9)
    hdr_cells[4].paragraphs[0].add_run('EURO').font.size = Pt(9)
    hdr_cells[4].paragraphs[0].runs[0].font.bold = True
    hdr_cells[4].paragraphs[0].runs[1].font.italic = True
    hdr_cells[4].paragraphs[0].runs[2].font.bold = True
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    hdr.height = Cm(1.25)
    hdr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    barra_cabeza = table.add_row()
    barra_cabeza_tabla = barra_cabeza.cells

    barra_cabeza_tabla[4].merge(barra_cabeza_tabla[3])
    barra_cabeza_tabla[3].merge(barra_cabeza_tabla[2])
    barra_cabeza_tabla[2].merge(barra_cabeza_tabla[1])
    barra_cabeza_tabla[1].merge(barra_cabeza_tabla[0])

    barra_cabeza.height = Cm(0.65)
    barra_cabeza.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    insert_hr(barra_cabeza_tabla[0].paragraphs[0])

    set_repeat_table_header(table.rows[0])
    set_repeat_table_header(table.rows[1])

    return table


def crear_tabla_resumen_pedido(pedido, table_resumen, icoterm, portes, transportista, peso, contacto, tel_fijo,
                               importe_bruto, imp_portes, dtopp, imp_dtopp, base_imp, iva, imp_iva, imp_rec_quiv,
                               rec_quiv, total, forma_pago, iban, giros, dp1, dp2, dp3):
    if str(icoterm).strip() != '':
        table_resumen.cell(0, 0).paragraphs[0].add_run('ICOTERM.').font.size = Pt(8)
        table_resumen.cell(0, 0).paragraphs[0].runs[0].font.bold = True
        table_resumen.cell(0, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(2, 0).paragraphs[0].add_run('PORTES').font.size = Pt(8)
    table_resumen.cell(2, 0).paragraphs[0].runs[0].font.bold = True
    table_resumen.cell(2, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(3, 0).paragraphs[0].add_run('TRANSPORT').font.size = Pt(8)
    table_resumen.cell(3, 0).paragraphs[0].runs[0].font.italic = True
    table_resumen.cell(3, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(5, 0).paragraphs[0].add_run('TRANSPORTE').font.size = Pt(8)
    table_resumen.cell(5, 0).paragraphs[0].runs[0].font.bold = True
    table_resumen.cell(5, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(6, 0).paragraphs[0].add_run('TRANSPORT').font.size = Pt(8)
    table_resumen.cell(6, 0).paragraphs[0].runs[0].font.italic = True
    table_resumen.cell(6, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(8, 0).paragraphs[0].add_run('PESO').font.size = Pt(8)
    table_resumen.cell(8, 0).paragraphs[0].runs[0].font.bold = True
    table_resumen.cell(8, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(9, 0).paragraphs[0].add_run('WEIGHT').font.size = Pt(8)
    table_resumen.cell(9, 0).paragraphs[0].runs[0].font.italic = True
    table_resumen.cell(9, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(11, 0).paragraphs[0].add_run('CONTACTO\n').font.size = Pt(8)
    table_resumen.cell(11, 0).paragraphs[0].runs[0].font.bold = True
    table_resumen.cell(11, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(12, 0).paragraphs[0].add_run('CONTACT PERSON').font.size = Pt(8)
    table_resumen.cell(12, 0).paragraphs[0].runs[0].font.italic = True
    table_resumen.cell(12, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(0, 1).paragraphs[0].add_run(icoterm).font.size = Pt(8)
    table_resumen.cell(0, 1).paragraphs[0].runs[0].font.italic = True

    if str(portes).strip() == 'D':
        table_resumen.cell(2, 1).paragraphs[0].add_run('Portes debidos')
        table_resumen.cell(3, 1).paragraphs[0].add_run('Transport not included')
    elif str(portes).strip() == 'P':
        table_resumen.cell(2, 1).paragraphs[0].add_run('Portes pagados')
        table_resumen.cell(3, 1).paragraphs[0].add_run('Transport included')
    elif str(portes).strip() == 'F':
        table_resumen.cell(2, 1).paragraphs[0].add_run('Portes en factura')
        table_resumen.cell(3, 1).paragraphs[0].add_run('Transport in invoice')
    else:
        table_resumen.cell(2, 1).paragraphs[0].text = 'A concretar'
        table_resumen.cell(3, 1).paragraphs[0].text = 'To be determined'

    table_resumen.cell(2, 1).paragraphs[0].runs[0].font.size = Pt(8)
    table_resumen.cell(3, 1).paragraphs[0].runs[0].font.size = Pt(8)
    table_resumen.cell(2, 1).paragraphs[0].runs[0].font.bold = True
    table_resumen.cell(3, 1).paragraphs[0].runs[0].font.italic = True

    if transportista.strip() == '':
        table_resumen.cell(5, 1).paragraphs[0].add_run('A concretar\n').font.size = Pt(8)
        table_resumen.cell(5, 1).paragraphs[0].runs[0].font.bold = True
        table_resumen.cell(6, 1).paragraphs[0].add_run('To be determined').font.size = Pt(8)
        table_resumen.cell(6, 1).paragraphs[0].runs[0].font.italic = True
    else:
        table_resumen.cell(5, 1).paragraphs[0].text = transportista
        table_resumen.cell(5, 1).paragraphs[0].runs[0].font.size = Pt(8)

    table_resumen.cell(8, 1).paragraphs[0].text = peso
    table_resumen.cell(8, 1).paragraphs[0].runs[0].font.size = Pt(8)

    if pedido:
        table_resumen.cell(11, 1).paragraphs[0].add_run(contacto)
    else:
        table_resumen.cell(11, 1).paragraphs[0].add_run(comprovar_usuario(contacto))

    table_resumen.cell(11, 1).paragraphs[0].runs[0].font.size = Pt(10)

    table_resumen.cell(12, 1).paragraphs[0].add_run(tel_fijo)
    table_resumen.cell(12, 1).paragraphs[0].runs[0].font.size = Pt(10)

    table_resumen.cell(1, 3).paragraphs[0].add_run('IMPORTE BRUTO / ').font.size = Pt(8)
    table_resumen.cell(1, 3).paragraphs[0].add_run('GROSS AMOUNT').font.size = Pt(8)
    table_resumen.cell(1, 3).paragraphs[0].runs[0].font.bold = True
    table_resumen.cell(1, 3).paragraphs[0].runs[1].font.italic = True
    table_resumen.cell(1, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(1, 5).paragraphs[0].add_run(importe_bruto + ' €').font.size = Pt(8)
    table_resumen.cell(1, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    table_resumen.cell(2, 3).paragraphs[0].add_run('PORTES / ').font.size = Pt(8)
    table_resumen.cell(2, 3).paragraphs[0].add_run('TRANSPORT').font.size = Pt(8)
    table_resumen.cell(2, 3).paragraphs[0].runs[0].font.bold = True
    table_resumen.cell(2, 3).paragraphs[0].runs[1].font.italic = True
    table_resumen.cell(2, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(2, 5).paragraphs[0].add_run(imp_portes + ' €').font.size = Pt(8)
    table_resumen.cell(2, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    table_resumen.cell(3, 3).paragraphs[0].add_run('DTOP. PP.').font.size = Pt(8)
    table_resumen.cell(3, 3).paragraphs[0].runs[0].font.bold = True
    table_resumen.cell(3, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(3, 4).paragraphs[0].add_run(dtopp + ' %').font.size = Pt(8)
    table_resumen.cell(3, 4).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    table_resumen.cell(3, 5).paragraphs[0].add_run(imp_dtopp + ' €').font.size = Pt(8)
    table_resumen.cell(3, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    table_resumen.cell(4, 3).paragraphs[0].add_run('BASE IMPONIBLE / ').font.size = Pt(8)
    table_resumen.cell(4, 3).paragraphs[0].add_run('TAXABLE BASE').font.size = Pt(8)
    table_resumen.cell(4, 3).paragraphs[0].runs[0].font.bold = True
    table_resumen.cell(4, 3).paragraphs[0].runs[1].font.italic = True
    table_resumen.cell(4, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(4, 5).paragraphs[0].add_run(base_imp + ' €').font.size = Pt(8)
    table_resumen.cell(4, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    table_resumen.cell(5, 3).paragraphs[0].add_run('IVA / ').font.size = Pt(8)
    table_resumen.cell(5, 3).paragraphs[0].add_run('IVA').font.size = Pt(8)
    table_resumen.cell(5, 3).paragraphs[0].runs[0].font.bold = True
    table_resumen.cell(5, 3).paragraphs[0].runs[1].font.italic = True
    table_resumen.cell(5, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(5, 4).paragraphs[0].add_run(iva + ' %').font.size = Pt(8)
    table_resumen.cell(5, 4).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    table_resumen.cell(5, 5).paragraphs[0].add_run(imp_iva + ' €').font.size = Pt(8)
    table_resumen.cell(5, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    table_resumen.cell(6, 3).paragraphs[0].add_run('REC. EQUIVALENCIA').font.size = Pt(8)
    table_resumen.cell(6, 3).paragraphs[0].runs[0].font.bold = True
    table_resumen.cell(6, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(6, 4).paragraphs[0].add_run(rec_quiv + ' %').font.size = Pt(8)
    table_resumen.cell(6, 4).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    table_resumen.cell(6, 5).paragraphs[0].add_run(imp_rec_quiv + ' €').font.size = Pt(8)
    table_resumen.cell(6, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    table_resumen.cell(7, 3).merge(table_resumen.cell(7, 4))
    table_resumen.cell(7, 4).merge(table_resumen.cell(7, 5))
    insert_hr(table_resumen.cell(7, 3).paragraphs[0])

    table_resumen.cell(8, 3).paragraphs[0].add_run('IMPORTE TOTAL / ').font.size = Pt(9)
    table_resumen.cell(8, 3).paragraphs[0].add_run('TOTAL AMOUNT').font.size = Pt(9)
    table_resumen.cell(8, 3).paragraphs[0].runs[0].font.bold = True
    table_resumen.cell(8, 3).paragraphs[0].runs[1].font.italic = True
    table_resumen.cell(8, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(8, 5).paragraphs[0].add_run(total + ' €').font.size = Pt(9)
    table_resumen.cell(8, 5).paragraphs[0].runs[0].font.bold = True
    table_resumen.cell(8, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    table_resumen.cell(11, 3).paragraphs[0].add_run('FORMA DE PAGO / ').font.size = Pt(8)
    table_resumen.cell(11, 3).paragraphs[0].add_run('MEANS OF PAYMENT').font.size = Pt(8)
    table_resumen.cell(11, 3).paragraphs[0].runs[0].font.bold = True
    table_resumen.cell(11, 3).paragraphs[0].runs[1].font.italic = True
    table_resumen.cell(11, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_resumen.cell(12, 5).merge(table_resumen.cell(12, 4))

    if str(forma_pago).strip() == 'TRANSFERENCIA' or str(forma_pago).strip() == 'CONTADO':
        table_resumen.cell(11, 5).paragraphs[0].add_run(str(forma_pago).strip()).font.size = Pt(8)
        table_resumen.cell(11, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        table_resumen.cell(12, 3).paragraphs[0].add_run('SWIFT/IBAN').font.size = Pt(9)
        table_resumen.cell(12, 3).paragraphs[0].runs[0].font.bold = True
        table_resumen.cell(12, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        table_resumen.cell(12, 5).paragraphs[0].add_run(iban).font.size = Pt(9)
        table_resumen.cell(12, 5).paragraphs[0].runs[0].font.bold = True
        table_resumen.cell(12, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif str(forma_pago).strip() == 'GIRO':
        table_resumen.cell(11, 5).paragraphs[0].add_run(
            str(forma_pago).strip() + ' a ' + giros + ' DIAS').font.size = Pt(8)
        table_resumen.cell(11, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        table_resumen.cell(12, 5).paragraphs[0].add_run('DIAS ').font.size = Pt(8)

        if str(dp1).strip() != '0' or str(dp1).strip() != '':
            table_resumen.cell(12, 5).paragraphs[0].add_run(dp1).font.size = Pt(8)
            table_resumen.cell(12, 5).paragraphs[0].runs[1].font.bold = True
            table_resumen.cell(12, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        if str(dp2).strip() != '0' or str(dp1).strip() != '':
            table_resumen.cell(12, 5).paragraphs[0].add_run('/' + dp2).font.size = Pt(8)
            table_resumen.cell(12, 5).paragraphs[0].runs[2].font.bold = True
            table_resumen.cell(12, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        if str(dp3).strip() != '0' or str(dp1).strip() != '':
            table_resumen.cell(12, 5).paragraphs[0].add_run('/' + dp3).font.size = Pt(8)
            table_resumen.cell(12, 5).paragraphs[0].runs[3].font.bold = True
            table_resumen.cell(12, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    else:
        table_resumen.cell(11, 5).paragraphs[0].add_run(str(forma_pago).strip()).font.size = Pt(8)
        table_resumen.cell(11, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def insertar_barra_final_productos(table):
    barra_pie = table.add_row()
    barra_pie_tabla = barra_pie.cells

    barra_pie_tabla[5].merge(barra_pie_tabla[4])
    barra_pie_tabla[4].merge(barra_pie_tabla[3])
    barra_pie_tabla[3].merge(barra_pie_tabla[2])
    barra_pie_tabla[2].merge(barra_pie_tabla[1])
    barra_pie_tabla[1].merge(barra_pie_tabla[0])

    barra_pie.height = Cm(0.2)
    barra_pie.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    barra_pie_tabla[0].paragraphs[0].add_run('.')
    barra_pie_tabla[0].paragraphs[0].runs[0].font.size = Pt(3)
    barra_pie_tabla[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    insert_hr(barra_pie_tabla[0].paragraphs[0])


def parse_floats(num):
    # Eliminar espacios en blanco y reemplazar puntos y comas
    num_parsed = str(num).strip().replace('.', '').replace(',', '.')

    # Convertir a float y formatear con dos decimales
    num_float = "{:.2f}".format(float(num_parsed))

    # Separar los miles con puntos y reemplazar el punto decimal por una coma
    parte_entera, parte_decimal = num_float.split('.')
    parte_entera_con_puntos = "{:,}".format(int(parte_entera)).replace(',', '.')
    num_parsed = parte_entera_con_puntos + ',' + parte_decimal

    return num_parsed
