import os

from django.shortcuts import render, redirect
from django.views.generic.base import View

import docx
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

from ofertador.forms import CargarOferta

import csv

from docxtpl import DocxTemplate


def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row


def insertHR(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(
        pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


def comprovar_plazo(fecha):
    array_fecha = fecha.split("/")

    if array_fecha[1] != '00' and array_fecha[2] != '0000':
        return fecha
    elif array_fecha[1] == '00' and array_fecha[2] == '0000':
        return str(array_fecha[0]) + ' dias'
    elif array_fecha[0] == '00' and array_fecha[2] == '0000':
        if array_fecha[1] == '01' or array_fecha[1] == '1':
            return str(array_fecha[1]) + ' mes'
        else:
            return str(array_fecha[1]) + ' meses'
    elif array_fecha[0] != '00' and array_fecha[1] != '00' and array_fecha[2] == '0000':
        if array_fecha[1] == '01' or array_fecha[1] == '1':
            return str(array_fecha[0]) + ' dias y ' + str(array_fecha[1]) + ' mes'
        else:
            return str(array_fecha[0]) + ' dias y ' + str(array_fecha[1]) + ' meses'
    elif array_fecha[0] == '00' and array_fecha[1] == '00' and array_fecha[2] != '0000':
        return str(array_fecha[2])


def comprovar_stock(fecha_pedido, fecha_plazo):
    if str(fecha_pedido) == str(fecha_plazo):
        return True
    else:
        array_fecha_pedido = str(fecha_pedido).split('/')
        array_fecha_plazo = str(fecha_plazo).split('/')

        if (array_fecha_pedido[0][0] == '0' or array_fecha_pedido[0][0] == ' ') and (
                array_fecha_plazo[0][0] == '0' or array_fecha_pedido[0][0] == ' '):
            if (array_fecha_pedido[1][0] == '0' or array_fecha_pedido[1][0] == ' ') and (
                    array_fecha_plazo[1][0] == '0' or array_fecha_pedido[1][0] == ' '):
                if array_fecha_pedido[0][1] == array_fecha_plazo[0][1] and array_fecha_pedido[1][1] == \
                        array_fecha_plazo[1][1] and array_fecha_pedido[2] == array_fecha_plazo[2]:
                    return True
                else:
                    return False
        else:
            if (array_fecha_pedido[1][0] == '0' or array_fecha_pedido[1][0] == ' ') and (
                    array_fecha_plazo[1][0] == '0' or array_fecha_pedido[1][0] == ' '):
                if array_fecha_pedido[0] == array_fecha_plazo[0] and array_fecha_pedido[1][1] == array_fecha_plazo[1][
                    1] and array_fecha_pedido[2] == array_fecha_plazo[2]:
                    return True
                else:
                    return False


class Ofertas(View):
    def get(self, request):
        form = CargarOferta()
        msg = ''
        titulo = 'Generador de ofertas'
        return render(request, 'index.html', {'form': form, 'mensaje': msg, 'titulo': titulo})

    def post(self, request):
        if request.POST:
            form = CargarOferta(request.POST, request.FILES)
            if form.is_valid():
                archivo_oferta = form.cleaned_data.get('oferta')

                with open('csvofertas/oferta.csv', 'wb+') as destination:
                    for chunk in archivo_oferta.chunks():
                        destination.write(chunk)

                oferta = ''
                fecha = ''
                validez = ''
                cliente = ''
                proveedor = ''
                rsoc = ''
                empresa = ''
                dir = ''
                cp = ''
                pob = ''
                pro = ''
                tel = ''
                mail = ''
                cab = ''
                peso = ''
                contacto = ''
                importe_bruto = ''
                portes = ''
                imp_portes = ''
                dtopp = ''
                imp_dtopp = ''
                base_imp = ''
                iva = ''
                imp_iva = ''
                rec_quiv = ''
                imp_rec_quiv = ''
                total = ''
                forma_pago = ''
                transportista = ''
                iban = 'ES25 2100-1083-1102-0005-4013'
                tel_fijo = '+34 937144561'

                doc = DocxTemplate("csvofertas/plantilla.docx")

                with open('csvofertas/oferta.csv') as csv_file:
                    csv_reader = csv.reader(csv_file, delimiter=';')
                    line_count = 0

                    for row in csv_reader:
                        if line_count == 1:
                            oferta = row[0]
                            fecha = row[1]
                            validez = row[2]
                            cliente = row[3]
                            proveedor = row[4]
                            rsoc = row[5]
                            empresa = row[6]
                            dir = row[7]
                            cp = row[8]
                            pob = row[9]
                            pro = row[10]
                            tel = row[11]
                            mail = row[13]
                            cab = str(row[42]).strip()
                            peso = row[17] + ' kg.'
                            contacto = row[32]
                            importe_bruto = row[20]
                            portes = row[18]
                            imp_portes = row[19]
                            dtopp = row[22]
                            imp_dtopp = row[23]
                            base_imp = row[26]
                            iva = row[27]
                            imp_iva = row[28]
                            rec_quiv = row[29]
                            imp_rec_quiv = row[30]
                            total = row[31]
                            forma_pago = row[49]
                            transportista = row[50]
                        line_count += 1

                context = \
                    {
                        'OFERTA': oferta,
                        'FECHA': fecha,
                        'VALIDEZ': validez,
                        'CLIENTE': cliente,
                        'PROVEEDOR': proveedor,
                        'RSOC': rsoc,
                        'EMPRESA': empresa,
                        'DIR': dir,
                        'CP': cp,
                        'POB': pob,
                        'PRO': pro,
                        'TEL': tel,
                        'MAIL': mail,
                        'CABECERA': cab,
                    }

                doc.render(context)
                nombre_oferta = str(archivo_oferta).split('.')[0]
                ruta_guardado = 'C:/generador/ofertas/' + nombre_oferta + '.docx'
                doc.save(ruta_guardado)

                doc = docx.Document(ruta_guardado)

                table = doc.add_table(rows=1, cols=6)

                table.allow_autofit = False

                for i in range(6):
                    for cell in table.columns[i].cells:
                        if i == 0:
                            cell.width = Cm(2.75)
                        elif i == 1:
                            cell.width = Cm(19)
                        elif i == 2:
                            cell.width = Cm(2.25)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        elif i == 3:
                            cell.width = Cm(2)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        elif i == 4:
                            cell.width = Cm(1.25)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        elif i == 5:
                            cell.width = Cm(2)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr = table.rows[0]
                hdr_cells = hdr.cells

                hdr_cells[0].paragraphs[0].add_run('\n\nREF.\n').font.size = Pt(9)
                hdr_cells[0].paragraphs[0].add_run('REF.\n').font.size = Pt(9)
                hdr_cells[0].paragraphs[0].runs[0].font.bold = True
                hdr_cells[0].paragraphs[0].runs[1].font.italic = True

                hdr_cells[1].paragraphs[0].add_run('\n\nDESCRIPCION\n').font.size = Pt(9)
                hdr_cells[1].paragraphs[0].add_run('SPECIFICATION\n').font.size = Pt(9)
                hdr_cells[1].paragraphs[0].runs[0].font.bold = True
                hdr_cells[1].paragraphs[0].runs[1].font.italic = True
                hdr_cells[1].paragraphs[0].runs[1].font.bold = False

                hdr_cells[2].paragraphs[0].add_run('\n\nCANTIDAD\n').font.size = Pt(9)
                hdr_cells[2].paragraphs[0].add_run('QUANTITY\n').font.size = Pt(9)
                hdr_cells[2].paragraphs[0].runs[0].font.bold = True
                hdr_cells[2].paragraphs[0].runs[1].font.italic = True
                hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr_cells[3].paragraphs[0].add_run('\n\nPRECIO\n').font.size = Pt(9)
                hdr_cells[3].paragraphs[0].add_run('PRICE\n').font.size = Pt(9)
                hdr_cells[3].paragraphs[0].add_run('EUROx100').font.size = Pt(9)
                hdr_cells[3].paragraphs[0].runs[0].font.bold = True
                hdr_cells[3].paragraphs[0].runs[1].font.italic = True
                hdr_cells[3].paragraphs[0].runs[2].font.bold = True
                hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr_cells[4].paragraphs[0].add_run('\n\nDTO.\n').font.size = Pt(9)
                hdr_cells[4].paragraphs[0].add_run('DIS.\n').font.size = Pt(9)
                hdr_cells[4].paragraphs[0].add_run('%').font.size = Pt(9)
                hdr_cells[4].paragraphs[0].runs[0].font.bold = True
                hdr_cells[4].paragraphs[0].runs[1].font.italic = True
                hdr_cells[4].paragraphs[0].runs[2].font.bold = True
                hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr_cells[5].paragraphs[0].add_run('\n\nIMPORTE\n').font.size = Pt(9)
                hdr_cells[5].paragraphs[0].add_run('AMOUNT\n').font.size = Pt(9)
                hdr_cells[5].paragraphs[0].add_run('EURO').font.size = Pt(9)
                hdr_cells[5].paragraphs[0].runs[0].font.bold = True
                hdr_cells[5].paragraphs[0].runs[1].font.italic = True
                hdr_cells[5].paragraphs[0].runs[2].font.bold = True
                hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr.height = Cm(2)
                hdr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                barra_cabeza = table.add_row()
                barra_cabeza_tabla = barra_cabeza.cells

                barra_cabeza_tabla[5].merge(barra_cabeza_tabla[4])
                barra_cabeza_tabla[4].merge(barra_cabeza_tabla[3])
                barra_cabeza_tabla[3].merge(barra_cabeza_tabla[2])
                barra_cabeza_tabla[2].merge(barra_cabeza_tabla[1])
                barra_cabeza_tabla[1].merge(barra_cabeza_tabla[0])

                barra_cabeza.height = Cm(0.65)
                barra_cabeza.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                insertHR(barra_cabeza_tabla[0].paragraphs[0])

                set_repeat_table_header(table.rows[0])
                set_repeat_table_header(table.rows[1])

                linias = 0

                with open('csvofertas/oferta.csv') as csv_file:
                    csv_reader = csv.reader(csv_file, delimiter=';')
                    count = 0

                    for row in csv_reader:
                        if count > 2:
                            if linias == 16 or (linias - 16) % 18 == 0:
                                if linias != 0:
                                    row_line = table.add_row()
                                    row_line_tabla = row_line.cells
                                    row_line_tabla[5].merge(row_line_tabla[4])
                                    row_line_tabla[4].merge(row_line_tabla[3])
                                    row_line_tabla[3].merge(row_line_tabla[2])
                                    row_line_tabla[2].merge(row_line_tabla[1])
                                    row_line_tabla[1].merge(row_line_tabla[0])

                                    row_line.height = Cm(0.65)
                                    row_line.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                                    insertHR(row_line_tabla[0].paragraphs[0])

                                    row_line = table.add_row()
                                    row_line_tabla = row_line.cells
                                    row_line_tabla[5].text = "Sigue..."
                                    row_line_tabla[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            row_prod = table.add_row()
                            row_cells = row_prod.cells

                            row_prod.height = Cm(1)
                            row_prod.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                            if str(row[23]).strip() == 'Texto':
                                row_cells[1].paragraphs[0].add_run(row[5]).font.size = Pt(8.5)
                            else:
                                row_cells[0].paragraphs[0].add_run(row[22]).font.size = Pt(10)
                                row_cells[0].paragraphs[0].add_run('\n' + row[4]).font.size = Pt(10)
                                row_cells[0].paragraphs[0].runs[1].font.italic = True

                                if comprovar_stock(str(fecha), str(row[16]).strip()):
                                    if str(row[23]).strip() == 'Especial':
                                        row_cells[1].paragraphs[0].add_run(row[5]).font.size = Pt(8.5)
                                    else:
                                        row_cells[1].paragraphs[0].add_run(row[23]).font.size = Pt(8.5)

                                    row_cells[1].paragraphs[0].add_run('\nPLAZO/').font.size = Pt(8)
                                    row_cells[1].paragraphs[0].add_run('Delivery:').font.size = Pt(8)
                                    row_cells[1].paragraphs[0].add_run('  [STOCK]').font.size = Pt(8)
                                    row_cells[1].paragraphs[0].runs[2].font.italic = True
                                    row_cells[1].paragraphs[0].runs[3].font.bold = True
                                else:
                                    if str(row[23]).strip() == 'Especial':
                                        row_cells[1].paragraphs[0].add_run(row[5]).font.size = Pt(8.5)
                                    else:
                                        row_cells[1].paragraphs[0].add_run(row[23]).font.size = Pt(8.5)

                                    row_cells[1].paragraphs[0].add_run('\nPLAZO/').font.size = Pt(8)
                                    row_cells[1].paragraphs[0].add_run('Delivery:').font.size = Pt(8)
                                    row_cells[1].paragraphs[0].add_run(
                                        '  ' + str(comprovar_plazo(row[16].strip()))).font.size = Pt(8)
                                    row_cells[1].paragraphs[0].runs[2].font.italic = True
                                    row_cells[1].paragraphs[0].runs[3].font.bold = True

                                row_cells[2].text = row[9]
                                row_cells[2].paragraphs[0].runs[0].font.size = Pt(10)
                                row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                                row_cells[3].text = row[18]
                                row_cells[3].paragraphs[0].runs[0].font.size = Pt(10)
                                row_cells[3].paragraphs[0].runs[0].font.bold = True
                                row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                                if str(row[19]).strip() == '':
                                    row_cells[4].text = 'NETO'
                                else:
                                    row_cells[4].text = row[19]

                                row_cells[4].paragraphs[0].runs[0].font.size = Pt(10)
                                row_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                                row_cells[5].text = row[20]
                                row_cells[5].paragraphs[0].runs[0].font.size = Pt(10)
                                row_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            linias += 1

                        count += 1

                barra_pie = table.add_row()
                barra_pie_tabla = barra_pie.cells

                barra_pie_tabla[5].merge(barra_pie_tabla[4])
                barra_pie_tabla[4].merge(barra_pie_tabla[3])
                barra_pie_tabla[3].merge(barra_pie_tabla[2])
                barra_pie_tabla[2].merge(barra_pie_tabla[1])
                barra_pie_tabla[1].merge(barra_pie_tabla[0])

                barra_pie.height = Cm(0.65)
                barra_pie.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                insertHR(barra_pie_tabla[0].paragraphs[0])

                pie_tabla = table.add_row().cells

                pie_tabla[5].merge(pie_tabla[4])
                pie_tabla[4].merge(pie_tabla[3])
                pie_tabla[3].merge(pie_tabla[2])
                pie_tabla[0].merge(pie_tabla[1])

                pie_tabla[0].text = 'PRECIOS VÁLIDOS PARA LAS CANTIDADES OFERTADAS'
                pie_tabla[0].paragraphs[0].runs[0].font.size = Pt(10)
                pie_tabla[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                pie_tabla[0].paragraphs[0].runs[0].font.bold = True

                pie_tabla[2].text = 'OFERTA VIGENTE HASTA LA FECHA:\n' + validez
                pie_tabla[2].paragraphs[0].runs[0].font.size = Pt(10)
                pie_tabla[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                pie_tabla[2].paragraphs[0].runs[0].font.bold = True

                if 16 >= linias > 4 or linias > 16 and (linias - 16) % 18 > 7:
                    doc.add_page_break()
                    doc.add_paragraph("\n\n\n")
                else:
                    doc.add_paragraph().add_run().font.size = Pt(6)

                table_resumen = doc.add_table(rows=12, cols=6)

                for i in range(5):
                    for cell in table_resumen.columns[i].cells:
                        if i == 0:
                            cell.width = Cm(2.49)
                        if i == 1:
                            cell.width = Cm(4.5)
                        if i == 2:
                            cell.width = Cm(1)
                        if i == 3:
                            cell.width = Cm(5.25)
                        if i == 4:
                            cell.width = Cm(1.73)
                        if i == 5:
                            cell.width = Cm(3.08)

                i = 0

                for row in table_resumen.rows:
                    if i == 6:
                        row.height = Cm(0.74)
                        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                    else:
                        row.height = Cm(0.35)
                        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                    i += 1

                '''Información del pedido'''

                table_resumen.cell(0, 0).paragraphs[0].add_run('PORTES').font.size = Pt(8)
                table_resumen.cell(0, 0).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(0, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(1, 0).paragraphs[0].add_run('TRANSPORT').font.size = Pt(8)
                table_resumen.cell(1, 0).paragraphs[0].runs[0].font.italic = True
                table_resumen.cell(1, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(3, 0).paragraphs[0].add_run('TRANSPORTE').font.size = Pt(8)
                table_resumen.cell(3, 0).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(3, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(4, 0).paragraphs[0].add_run('TRANSPORT').font.size = Pt(8)
                table_resumen.cell(4, 0).paragraphs[0].runs[0].font.italic = True
                table_resumen.cell(4, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(6, 0).paragraphs[0].add_run('PESO\n').font.size = Pt(8)
                table_resumen.cell(6, 0).paragraphs[0].add_run('WEIGHT').font.size = Pt(8)
                table_resumen.cell(6, 0).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(6, 0).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(6, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(8, 0).paragraphs[0].add_run('CONTACTO\n').font.size = Pt(8)
                table_resumen.cell(8, 0).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(8, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(9, 0).paragraphs[0].add_run('CONTACT PERSON').font.size = Pt(8)
                table_resumen.cell(9, 0).paragraphs[0].runs[0].font.italic = True
                table_resumen.cell(9, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                if str(portes).strip() == 'D':
                    table_resumen.cell(0, 1).paragraphs[0].add_run('Portes debidos')
                    table_resumen.cell(1, 1).paragraphs[0].add_run('Transport not included')
                elif str(portes).strip() == 'P':
                    table_resumen.cell(0, 1).paragraphs[0].add_run('Portes pagados')
                    table_resumen.cell(1, 1).paragraphs[0].add_run('Transport included')
                elif str(portes).strip() == 'F':
                    table_resumen.cell(0, 1).paragraphs[0].add_run('Portes en factura')
                    table_resumen.cell(1, 1).paragraphs[0].add_run('Transport in invoice')
                else:
                    table_resumen.cell(0, 1).paragraphs[0].text = 'A concretar'
                    table_resumen.cell(1, 1).paragraphs[0].text = 'To be determined'

                table_resumen.cell(0, 1).paragraphs[0].runs[0].font.size = Pt(8)
                table_resumen.cell(1, 1).paragraphs[0].runs[0].font.size = Pt(8)
                table_resumen.cell(0, 1).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(1, 1).paragraphs[0].runs[0].font.italic = True

                if transportista.strip() == '':
                    table_resumen.cell(3, 1).paragraphs[0].add_run('A concretar\n').font.size = Pt(8)
                    table_resumen.cell(3, 1).paragraphs[0].add_run('To be determined').font.size = Pt(8)
                    table_resumen.cell(3, 1).paragraphs[0].runs[0].font.bold = True
                    table_resumen.cell(3, 1).paragraphs[0].runs[1].font.italic = True
                else:
                    table_resumen.cell(3, 1).paragraphs[0].text = transportista
                    table_resumen.cell(3, 1).paragraphs[0].runs[0].font.size = Pt(8)

                table_resumen.cell(6, 1).paragraphs[0].text = peso
                table_resumen.cell(6, 1).paragraphs[0].runs[0].font.size = Pt(8)

                table_resumen.cell(8, 1).paragraphs[0].add_run(contacto + '\n')
                table_resumen.cell(8, 1).paragraphs[0].runs[0].font.size = Pt(10)
                table_resumen.cell(8, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                table_resumen.cell(9, 1).paragraphs[0].add_run(tel_fijo)
                table_resumen.cell(9, 1).paragraphs[0].runs[0].font.size = Pt(10)

                '''Resumen del pedido'''

                table_resumen.cell(0, 3).paragraphs[0].add_run('IMPORTE BRUTO / ').font.size = Pt(8)
                table_resumen.cell(0, 3).paragraphs[0].add_run('GROSS AMOUNT').font.size = Pt(8)
                table_resumen.cell(0, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(0, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(0, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(0, 5).paragraphs[0].add_run(importe_bruto + ' €').font.size = Pt(8)
                table_resumen.cell(0, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(1, 3).paragraphs[0].add_run('PORTES / ').font.size = Pt(8)
                table_resumen.cell(1, 3).paragraphs[0].add_run('TRANSPORT').font.size = Pt(8)
                table_resumen.cell(1, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(1, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(1, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(1, 5).paragraphs[0].add_run(imp_portes + ' €').font.size = Pt(8)
                table_resumen.cell(1, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(2, 3).paragraphs[0].add_run('DTOP. PP.').font.size = Pt(8)
                table_resumen.cell(2, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(2, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(2, 4).paragraphs[0].add_run(dtopp + ' %').font.size = Pt(8)
                table_resumen.cell(2, 4).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(2, 5).paragraphs[0].add_run(imp_dtopp + ' €').font.size = Pt(8)
                table_resumen.cell(2, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(3, 3).paragraphs[0].add_run('BASE IMPONIBLE / ').font.size = Pt(8)
                table_resumen.cell(3, 3).paragraphs[0].add_run('TAXABLE BASE').font.size = Pt(8)
                table_resumen.cell(3, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(3, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(3, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(3, 5).paragraphs[0].add_run(base_imp + ' €').font.size = Pt(8)
                table_resumen.cell(3, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(4, 3).paragraphs[0].add_run('IVA / ').font.size = Pt(8)
                table_resumen.cell(4, 3).paragraphs[0].add_run('IVA').font.size = Pt(8)
                table_resumen.cell(4, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(4, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(4, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(4, 4).paragraphs[0].add_run(iva + ' %').font.size = Pt(8)
                table_resumen.cell(4, 4).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(4, 5).paragraphs[0].add_run(imp_iva + ' €').font.size = Pt(8)
                table_resumen.cell(4, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(5, 3).paragraphs[0].add_run('REC. EQUIVALENCIA').font.size = Pt(8)
                table_resumen.cell(5, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(5, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(5, 4).paragraphs[0].add_run(rec_quiv + ' %').font.size = Pt(8)
                table_resumen.cell(5, 4).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(5, 5).paragraphs[0].add_run(imp_rec_quiv + ' €').font.size = Pt(8)
                table_resumen.cell(5, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(6, 3).merge(table_resumen.cell(6, 4))
                table_resumen.cell(6, 4).merge(table_resumen.cell(6, 5))
                insertHR(table_resumen.cell(6, 3).paragraphs[0])

                table_resumen.cell(7, 3).paragraphs[0].add_run('IMPORTE TOTAL / ').font.size = Pt(9)
                table_resumen.cell(7, 3).paragraphs[0].add_run('TOTAL AMOUNT').font.size = Pt(9)
                table_resumen.cell(7, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(7, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(7, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(7, 5).paragraphs[0].add_run(total + ' €').font.size = Pt(9)
                table_resumen.cell(7, 5).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(7, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(10, 3).paragraphs[0].add_run('FORMA DE PAGO / ').font.size = Pt(8)
                table_resumen.cell(10, 3).paragraphs[0].add_run('MEANS OF PAYMENT').font.size = Pt(8)
                table_resumen.cell(10, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(10, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(10, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(11, 5).merge(table_resumen.cell(11, 4))
                table_resumen.cell(11, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                if str(forma_pago).strip() == 'TRANSFERENCIA' or str(forma_pago).strip() == 'CONTADO':
                    table_resumen.cell(10, 5).paragraphs[0].add_run(str(forma_pago).strip()).font.size = Pt(8)
                    table_resumen.cell(10, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    table_resumen.cell(11, 3).paragraphs[0].add_run('SWIFT/IBAN').font.size = Pt(9)
                    table_resumen.cell(11, 3).paragraphs[0].runs[0].font.bold = True
                    table_resumen.cell(11, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                    table_resumen.cell(11, 5).paragraphs[0].add_run(iban).font.size = Pt(9)
                    table_resumen.cell(11, 5).paragraphs[0].runs[0].font.bold = True

                else:
                    table_resumen.cell(10, 5).paragraphs[0].add_run(str(forma_pago).strip()).font.size = Pt(8)
                    table_resumen.cell(10, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                barra_cond = doc.add_paragraph()
                insertHR(barra_cond)
                doc.add_paragraph("")

                condiciones = doc.add_paragraph()
                condiciones.add_run('CONDICIONES:\n').font.size = Pt(11)
                condiciones.add_run('\n').font.size = Pt(3)
                condiciones.add_run('- Disponibilidad y precios indicados salvo venta.\n- Estos ').font.size = Pt(9)
                condiciones.add_run('precios ').font.size = Pt(9)
                condiciones.add_run('son para la ').font.size = Pt(9)
                condiciones.add_run('totalidad de la oferta').font.size = Pt(9)
                condiciones.add_run(
                    ', en caso de pedido parcial los precios estarían sujetos a revisión.\n- Los ').font.size = Pt(9)
                condiciones.add_run('plazos de entrega ').font.size = Pt(9)
                condiciones.add_run('indicados son orientativos y se consideran, ').font.size = Pt(9)
                condiciones.add_run('días laborales y en nuestro almacén, ').font.size = Pt(9)
                condiciones.add_run('\n  a partir de la ').font.size = Pt(9)
                condiciones.add_run('fecha confirmación del pedido.\n').font.size = Pt(9)
                condiciones.add_run('- No se aceptan devolución de piezas especiales ').font.size = Pt(9)
                condiciones.add_run('ni medidas fuera de catálogo.\n').font.size = Pt(9)
                condiciones.add_run('- ').font.size = Pt(9)
                condiciones.add_run(
                    'Las piezas especiales se podrán suministrar con un +/- 10% de la cantidad ofertada.').font.size = Pt(
                    9)
                condiciones.add_run(
                    '\n\n- El suministro quedará supeditado a la concesión de riesgo por parte de ').font.size = Pt(9)
                condiciones.add_run('Crédito y Caución.').font.size = Pt(9)

                condiciones.runs[0].font.bold = True
                condiciones.runs[3].font.bold = True
                condiciones.runs[5].font.bold = True
                condiciones.runs[7].font.bold = True
                condiciones.runs[9].font.bold = True
                condiciones.runs[11].font.bold = True
                condiciones.runs[13].font.bold = True
                condiciones.runs[15].font.bold = True
                condiciones.runs[17].font.bold = True

                doc.save(ruta_guardado)
                # os.startfile(ruta_guardado, 'open')

                return redirect('inicio')
            else:
                form = CargarOferta()
                msg = 'Fichero no válido. Porfavor, compruebe el archivo.'
                titulo = 'Generador de ofertas'
                return render(request, 'index.html', {'form': form, 'mensaje': msg, 'titulo': titulo})


class Pedidos(View):
    def get(self, request):
        form = CargarOferta()
        msg = ''
        titulo = 'Generador de Pedidos'
        return render(request, 'index.html', {'form': form, 'mensaje': msg, 'titulo': titulo})

    def post(self, request):
        if request.POST:
            form = CargarOferta(request.POST, request.FILES)
            if form.is_valid():
                archivo_pedido = form.cleaned_data.get('oferta')

                with open('csvofertas/pedido.csv', 'wb+') as destination:
                    for chunk in archivo_pedido.chunks():
                        destination.write(chunk)

                pedido = ''
                fecha = ''
                plazo = ''
                cliente = ''
                proveedor = ''
                rsoc = ''
                empresa = ''
                dir = ''
                cp = ''
                pob = ''
                pro = ''
                nif = ''
                agente = ''
                tel = ''
                mail = ''
                peso = ''
                contacto = ''
                importe_bruto = ''
                portes = ''
                imp_portes = ''
                dtopp = ''
                imp_dtopp = ''
                base_imp = ''
                iva = ''
                imp_iva = ''
                rec_quiv = ''
                imp_rec_quiv = ''
                total = ''
                forma_pago = ''
                giros = ''
                dp1 = ''
                dp2 = ''
                dp3 = ''
                transportista = ''
                icoterm = ''
                iban = 'ES25 2100-1083-1102-0005-4013'
                tel_fijo = '+34 937144561'

                doc = DocxTemplate("csvofertas/plantilla_pedido.docx")

                with open('csvofertas/pedido.csv') as csv_file:
                    csv_reader = csv.reader(csv_file, delimiter=';')
                    line_count = 0

                    for row in csv_reader:
                        if line_count == 1:
                            pedido = row[0]
                            fecha = row[1]
                            plazo = row[2]
                            cliente = row[3]
                            proveedor = row[4]
                            rsoc = row[5]
                            empresa = row[6]
                            dir = row[7]
                            cp = row[8]
                            pob = row[9]
                            pro = row[10]
                            tel = row[11]
                            mail = row[13]
                            nif = row[15]
                            icoterm = row[17]
                            agente = row[18]
                            forma_pago = row[21]
                            giros = row[22]
                            dp1 = row[25]
                            dp2 = row[26]
                            dp3 = row[27]
                            contacto = row[28]
                            transportista = row[30]
                            peso = row[31] + ' kg.'
                            importe_bruto = row[32]
                            portes = row[33]
                            imp_portes = row[34]
                            dtopp = row[36]
                            imp_dtopp = row[37]
                            base_imp = row[40]
                            iva = row[41]
                            imp_iva = row[42]
                            rec_quiv = row[43]
                            imp_rec_quiv = row[44]
                            total = row[45]
                        line_count += 1

                context = \
                    {
                        'PEDIDO': pedido,
                        'FECHA': fecha,
                        'CLIENTE': cliente,
                        'PROVEEDOR': proveedor,
                        'AGENTE': agente,
                        'PLAZOENTREGA': plazo,
                        'RSOC': rsoc,
                        'EMPRESA': empresa,
                        'DIR': dir,
                        'CP': cp,
                        'POB': pob,
                        'PRO': pro,
                        'NIF': nif,
                        'TEL': tel,
                        'MAIL': mail,
                    }

                doc.render(context)
                nombre_pedido = str(archivo_pedido).split('.')[0]
                ruta_guardado = 'C:/generador/pedidos/' + nombre_pedido + '.docx'
                doc.save(ruta_guardado)

                doc = docx.Document(ruta_guardado)

                table = doc.add_table(rows=1, cols=6)

                for i in range(6):
                    for cell in table.columns[i].cells:
                        if i == 0:
                            cell.width = Cm(2.68)
                        elif i == 1:
                            cell.width = Cm(19)
                        elif i == 2:
                            cell.width = Inches(0.2)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        elif i == 3:
                            cell.width = Inches(0.2)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        elif i == 4:
                            cell.width = Inches(0.03)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        elif i == 5:
                            cell.width = Inches(0.2)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr = table.rows[0]
                hdr_cells = hdr.cells

                hdr_cells[0].paragraphs[0].add_run('REF.\n').font.size = Pt(9)
                hdr_cells[0].paragraphs[0].add_run('REF.\n').font.size = Pt(9)
                hdr_cells[0].paragraphs[0].runs[0].font.bold = True
                hdr_cells[0].paragraphs[0].runs[1].font.italic = True

                hdr_cells[1].paragraphs[0].add_run('DESCRIPCION\n').font.size = Pt(9)
                hdr_cells[1].paragraphs[0].add_run('SPECIFICATION\n').font.size = Pt(9)
                hdr_cells[1].paragraphs[0].runs[0].font.bold = True
                hdr_cells[1].paragraphs[0].runs[1].font.italic = True
                hdr_cells[1].paragraphs[0].runs[1].font.bold = False

                hdr_cells[2].paragraphs[0].add_run('CANTIDAD\n').font.size = Pt(9)
                hdr_cells[2].paragraphs[0].add_run('QUANTITY\n').font.size = Pt(9)
                hdr_cells[2].paragraphs[0].runs[0].font.bold = True
                hdr_cells[2].paragraphs[0].runs[1].font.italic = True
                hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr_cells[3].paragraphs[0].add_run('PRECIO\n').font.size = Pt(9)
                hdr_cells[3].paragraphs[0].add_run('PRICE\n').font.size = Pt(9)
                hdr_cells[3].paragraphs[0].add_run('EUROx100').font.size = Pt(9)
                hdr_cells[3].paragraphs[0].runs[0].font.bold = True
                hdr_cells[3].paragraphs[0].runs[1].font.italic = True
                hdr_cells[3].paragraphs[0].runs[2].font.bold = True
                hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr_cells[4].paragraphs[0].add_run('DTO.\n').font.size = Pt(9)
                hdr_cells[4].paragraphs[0].add_run('DIS.\n').font.size = Pt(9)
                hdr_cells[4].paragraphs[0].add_run('%').font.size = Pt(9)
                hdr_cells[4].paragraphs[0].runs[0].font.bold = True
                hdr_cells[4].paragraphs[0].runs[1].font.italic = True
                hdr_cells[4].paragraphs[0].runs[2].font.bold = True
                hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr_cells[5].paragraphs[0].add_run('IMPORTE\n').font.size = Pt(9)
                hdr_cells[5].paragraphs[0].add_run('AMOUNT\n').font.size = Pt(9)
                hdr_cells[5].paragraphs[0].add_run('EURO').font.size = Pt(9)
                hdr_cells[5].paragraphs[0].runs[0].font.bold = True
                hdr_cells[5].paragraphs[0].runs[1].font.italic = True
                hdr_cells[5].paragraphs[0].runs[2].font.bold = True
                hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr.height = Cm(1.25)
                hdr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                barra_cabeza = table.add_row()
                barra_cabeza_tabla = barra_cabeza.cells

                barra_cabeza_tabla[5].merge(barra_cabeza_tabla[4])
                barra_cabeza_tabla[4].merge(barra_cabeza_tabla[3])
                barra_cabeza_tabla[3].merge(barra_cabeza_tabla[2])
                barra_cabeza_tabla[2].merge(barra_cabeza_tabla[1])
                barra_cabeza_tabla[1].merge(barra_cabeza_tabla[0])

                barra_cabeza.height = Cm(0.65)
                barra_cabeza.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                insertHR(barra_cabeza_tabla[0].paragraphs[0])

                set_repeat_table_header(table.rows[0])
                set_repeat_table_header(table.rows[1])

                linias = 0

                with open('csvofertas/pedido.csv') as csv_file:
                    csv_reader = csv.reader(csv_file, delimiter=';')
                    count = 0

                    for row in csv_reader:
                        if count > 2:
                            if linias == 14 or (linias - 14) % 16 == 0:
                                if linias != 0:
                                    row_line = table.add_row()
                                    row_line_tabla = row_line.cells
                                    row_line_tabla[5].merge(row_line_tabla[4])
                                    row_line_tabla[4].merge(row_line_tabla[3])
                                    row_line_tabla[3].merge(row_line_tabla[2])
                                    row_line_tabla[2].merge(row_line_tabla[1])
                                    row_line_tabla[1].merge(row_line_tabla[0])

                                    row_line.height = Cm(0.65)
                                    row_line.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                                    insertHR(row_line_tabla[0].paragraphs[0])

                                    row_line = table.add_row()
                                    row_line_tabla = row_line.cells
                                    row_line_tabla[5].text = "Sigue..."
                                    row_line_tabla[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            row_prod = table.add_row()
                            row_cells = row_prod.cells

                            row_prod.height = Cm(1)
                            row_prod.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                            row_cells[0].paragraphs[0].add_run(row[8]).font.size = Pt(10)
                            row_cells[0].paragraphs[0].add_run('\n' + row[1]).font.size = Pt(10)
                            row_cells[0].paragraphs[0].runs[1].font.italic = True

                            row_cells[1].paragraphs[0].add_run(row[2]).font.size = Pt(8.5)
                            row_cells[1].paragraphs[0].add_run('\nPedido: ').font.size = Pt(8)
                            row_cells[1].paragraphs[0].add_run(row[0]).font.size = Pt(8)
                            row_cells[1].paragraphs[0].runs[2].font.bold = True

                            row_cells[2].text = row[3]
                            row_cells[2].paragraphs[0].runs[0].font.size = Pt(10)
                            row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            row_cells[3].text = row[4]
                            row_cells[3].paragraphs[0].runs[0].font.size = Pt(10)
                            row_cells[3].paragraphs[0].runs[0].font.bold = True
                            row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            row_cells[4].text = row[6]
                            row_cells[4].paragraphs[0].runs[0].font.size = Pt(10)
                            row_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            row_cells[5].text = row[7]
                            row_cells[5].paragraphs[0].runs[0].font.size = Pt(10)
                            row_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            linias += 1

                        count += 1

                barra_pie = table.add_row()
                barra_pie_tabla = barra_pie.cells

                barra_pie_tabla[5].merge(barra_pie_tabla[4])
                barra_pie_tabla[4].merge(barra_pie_tabla[3])
                barra_pie_tabla[3].merge(barra_pie_tabla[2])
                barra_pie_tabla[2].merge(barra_pie_tabla[1])
                barra_pie_tabla[1].merge(barra_pie_tabla[0])

                barra_pie.height = Cm(0.65)
                barra_pie.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                insertHR(barra_pie_tabla[0].paragraphs[0])

                pie_tabla = table.add_row().cells

                pie_tabla[5].merge(pie_tabla[4])
                pie_tabla[4].merge(pie_tabla[3])
                pie_tabla[3].merge(pie_tabla[2])
                pie_tabla[0].merge(pie_tabla[1])

                pie_tabla[2].paragraphs[0].add_run('PLAZO DE ENTREGA:\n').font.size = Pt(10)
                pie_tabla[2].paragraphs[0].add_run(plazo).font.size = Pt(10)
                pie_tabla[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                pie_tabla[2].paragraphs[0].runs[0].font.bold = True
                pie_tabla[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(191, 34, 34)
                pie_tabla[2].paragraphs[0].runs[1].font.bold = True
                pie_tabla[2].paragraphs[0].runs[1].font.italic = True

                doc.add_paragraph()

                print("Linias: " + str(linias))
                print((linias - 14) % 21)

                if 14 >= linias > 8 or linias > 14 and (linias - 14) % 16 > 8:
                    doc.add_page_break()
                    doc.add_paragraph("\n\n\n")

                table_resumen = doc.add_table(rows=13, cols=6)

                for i in range(5):
                    for cell in table_resumen.columns[i].cells:
                        if i == 0:
                            cell.width = Cm(2.49)
                        if i == 1:
                            cell.width = Cm(4.5)
                        if i == 2:
                            cell.width = Cm(1)
                        if i == 3:
                            cell.width = Cm(5.25)
                        if i == 4:
                            cell.width = Cm(1.73)
                        if i == 5:
                            cell.width = Cm(3.08)

                i = 0

                for row in table_resumen.rows:
                    if i == 7:
                        row.height = Cm(0.74)
                        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                    else:
                        row.height = Cm(0.35)
                        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                    i += 1

                '''Información del pedido'''

                if str(icoterm).strip() != '':
                    table_resumen.cell(0, 0).paragraphs[0].add_run('ICOTERM.').font.size = Pt(8)
                    table_resumen.cell(0, 0).paragraphs[0].runs[0].font.bold = True
                    table_resumen.cell(0, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(2, 0).paragraphs[0].add_run('PORTES').font.size = Pt(8)
                table_resumen.cell(2, 0).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(2, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(3, 0).paragraphs[0].add_run('TRANSPORT').font.size = Pt(8)
                table_resumen.cell(3, 0).paragraphs[0].runs[0].font.italic = True
                table_resumen.cell(3, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(5, 0).paragraphs[0].add_run('TRANSPORTE').font.size = Pt(8)
                table_resumen.cell(5, 0).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(5, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(6, 0).paragraphs[0].add_run('TRANSPORT').font.size = Pt(8)
                table_resumen.cell(6, 0).paragraphs[0].runs[0].font.italic = True
                table_resumen.cell(6, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(8, 0).paragraphs[0].add_run('PESO').font.size = Pt(8)
                table_resumen.cell(8, 0).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(8, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(9, 0).paragraphs[0].add_run('WEIGHT').font.size = Pt(8)
                table_resumen.cell(9, 0).paragraphs[0].runs[0].font.italic = True
                table_resumen.cell(9, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(11, 0).paragraphs[0].add_run('CONTACTO\n').font.size = Pt(8)
                table_resumen.cell(11, 0).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(11, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(12, 0).paragraphs[0].add_run('CONTACT PERSON').font.size = Pt(8)
                table_resumen.cell(12, 0).paragraphs[0].runs[0].font.italic = True
                table_resumen.cell(12, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(0, 1).paragraphs[0].add_run(icoterm).font.size = Pt(8)
                table_resumen.cell(0, 1).paragraphs[0].runs[0].font.italic = True

                if str(portes).strip() == 'D':
                    table_resumen.cell(2, 1).paragraphs[0].add_run('Portes debidos')
                    table_resumen.cell(3, 1).paragraphs[0].add_run('Transport not included')
                elif str(portes).strip() == 'P':
                    table_resumen.cell(2, 1).paragraphs[0].add_run('Portes pagados')
                    table_resumen.cell(3, 1).paragraphs[0].add_run('Transport included')
                elif str(portes).strip() == 'F':
                    table_resumen.cell(2, 1).paragraphs[0].add_run('Portes en factura')
                    table_resumen.cell(3, 1).paragraphs[0].add_run('Transport in invoice')
                else:
                    table_resumen.cell(2, 1).paragraphs[0].text = 'A concretar'
                    table_resumen.cell(3, 1).paragraphs[0].text = 'To be determined'

                table_resumen.cell(2, 1).paragraphs[0].runs[0].font.size = Pt(8)
                table_resumen.cell(3, 1).paragraphs[0].runs[0].font.size = Pt(8)
                table_resumen.cell(2, 1).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(3, 1).paragraphs[0].runs[0].font.italic = True

                if transportista.strip() == '':
                    table_resumen.cell(5, 1).paragraphs[0].add_run('A concretar\n').font.size = Pt(8)
                    table_resumen.cell(5, 1).paragraphs[0].runs[0].font.bold = True
                    table_resumen.cell(6, 1).paragraphs[0].add_run('To be determined').font.size = Pt(8)
                    table_resumen.cell(6, 1).paragraphs[0].runs[0].font.italic = True
                else:
                    table_resumen.cell(5, 1).paragraphs[0].text = transportista
                    table_resumen.cell(5, 1).paragraphs[0].runs[0].font.size = Pt(8)

                table_resumen.cell(8, 1).paragraphs[0].text = peso
                table_resumen.cell(8, 1).paragraphs[0].runs[0].font.size = Pt(8)

                table_resumen.cell(11, 1).paragraphs[0].add_run('          ' + contacto)
                table_resumen.cell(11, 1).paragraphs[0].runs[0].font.size = Pt(10)

                table_resumen.cell(12, 1).paragraphs[0].add_run(tel_fijo)
                table_resumen.cell(12, 1).paragraphs[0].runs[0].font.size = Pt(10)

                # Resumen del pedido

                table_resumen.cell(1, 3).paragraphs[0].add_run('IMPORTE BRUTO / ').font.size = Pt(8)
                table_resumen.cell(1, 3).paragraphs[0].add_run('GROSS AMOUNT').font.size = Pt(8)
                table_resumen.cell(1, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(1, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(1, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(1, 5).paragraphs[0].add_run(importe_bruto + ' €').font.size = Pt(8)
                table_resumen.cell(1, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(2, 3).paragraphs[0].add_run('PORTES / ').font.size = Pt(8)
                table_resumen.cell(2, 3).paragraphs[0].add_run('TRANSPORT').font.size = Pt(8)
                table_resumen.cell(2, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(2, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(2, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(2, 5).paragraphs[0].add_run(imp_portes + ' €').font.size = Pt(8)
                table_resumen.cell(2, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(3, 3).paragraphs[0].add_run('DTOP. PP.').font.size = Pt(8)
                table_resumen.cell(3, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(3, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(3, 4).paragraphs[0].add_run(dtopp + ' %').font.size = Pt(8)
                table_resumen.cell(3, 4).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(3, 5).paragraphs[0].add_run(imp_dtopp + ' €').font.size = Pt(8)
                table_resumen.cell(3, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(4, 3).paragraphs[0].add_run('BASE IMPONIBLE / ').font.size = Pt(8)
                table_resumen.cell(4, 3).paragraphs[0].add_run('TAXABLE BASE').font.size = Pt(8)
                table_resumen.cell(4, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(4, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(4, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(4, 5).paragraphs[0].add_run(base_imp + ' €').font.size = Pt(8)
                table_resumen.cell(4, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(5, 3).paragraphs[0].add_run('IVA / ').font.size = Pt(8)
                table_resumen.cell(5, 3).paragraphs[0].add_run('IVA').font.size = Pt(8)
                table_resumen.cell(5, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(5, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(5, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(5, 4).paragraphs[0].add_run(iva + ' %').font.size = Pt(8)
                table_resumen.cell(5, 4).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(5, 5).paragraphs[0].add_run(imp_iva + ' €').font.size = Pt(8)
                table_resumen.cell(5, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(6, 3).paragraphs[0].add_run('REC. EQUIVALENCIA').font.size = Pt(8)
                table_resumen.cell(6, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(6, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(6, 4).paragraphs[0].add_run(rec_quiv + ' %').font.size = Pt(8)
                table_resumen.cell(6, 4).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(6, 5).paragraphs[0].add_run(imp_rec_quiv + ' €').font.size = Pt(8)
                table_resumen.cell(6, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(7, 3).merge(table_resumen.cell(7, 4))
                table_resumen.cell(7, 4).merge(table_resumen.cell(7, 5))
                insertHR(table_resumen.cell(7, 3).paragraphs[0])

                table_resumen.cell(8, 3).paragraphs[0].add_run('IMPORTE TOTAL / ').font.size = Pt(9)
                table_resumen.cell(8, 3).paragraphs[0].add_run('TOTAL AMOUNT').font.size = Pt(9)
                table_resumen.cell(8, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(8, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(8, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(8, 5).paragraphs[0].add_run(total + ' €').font.size = Pt(9)
                table_resumen.cell(8, 5).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(8, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(11, 3).paragraphs[0].add_run('FORMA DE PAGO / ').font.size = Pt(8)
                table_resumen.cell(11, 3).paragraphs[0].add_run('MEANS OF PAYMENT').font.size = Pt(8)
                table_resumen.cell(11, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(11, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(11, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(12, 5).merge(table_resumen.cell(12, 4))

                if str(forma_pago).strip() == 'TRANSFERENCIA' or str(forma_pago).strip() == 'CONTADO':
                    table_resumen.cell(11, 5).paragraphs[0].add_run(str(forma_pago).strip()).font.size = Pt(8)
                    table_resumen.cell(11, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    table_resumen.cell(12, 3).paragraphs[0].add_run('SWIFT/IBAN').font.size = Pt(9)
                    table_resumen.cell(12, 3).paragraphs[0].runs[0].font.bold = True
                    table_resumen.cell(12, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                    table_resumen.cell(12, 5).paragraphs[0].add_run(iban).font.size = Pt(9)
                    table_resumen.cell(12, 5).paragraphs[0].runs[0].font.bold = True
                    table_resumen.cell(12, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif str(forma_pago).strip() == 'GIRO':
                    table_resumen.cell(11, 5).paragraphs[0].add_run(
                        str(forma_pago).strip() + ' a ' + giros + ' DIAS').font.size = Pt(8)
                    table_resumen.cell(12, 5).paragraphs[0].add_run('DIAS ').font.size = Pt(8)

                    if str(dp1).strip() != '0' or str(dp1).strip() != '':
                        table_resumen.cell(12, 5).paragraphs[0].add_run(dp1).font.size = Pt(8)
                        table_resumen.cell(12, 5).paragraphs[0].runs[1].font.bold = True
                        table_resumen.cell(12, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    if str(dp2).strip() != '0' or str(dp1).strip() != '':
                        table_resumen.cell(12, 5).paragraphs[0].add_run('/' + dp2).font.size = Pt(8)
                        table_resumen.cell(12, 5).paragraphs[0].runs[2].font.bold = True
                        table_resumen.cell(12, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    if str(dp3).strip() != '0' or str(dp1).strip() != '':
                        table_resumen.cell(12, 5).paragraphs[0].add_run('/' + dp3).font.size = Pt(8)
                        table_resumen.cell(12, 5).paragraphs[0].runs[3].font.bold = True
                        table_resumen.cell(12, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                else:
                    table_resumen.cell(11, 5).paragraphs[0].add_run(str(forma_pago).strip()).font.size = Pt(8)
                    table_resumen.cell(11, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                doc.save(ruta_guardado)

                # os.startfile(ruta_guardado)

                return redirect('pedidos')
            else:
                form = CargarOferta()
                msg = 'Fichero no válido. Porfavor, compruebe el archivo.'
                titulo = 'Generador de Pedidos'
                return render(request, 'index.html', {'form': form, 'mensaje': msg, 'titulo': titulo})


class PreAlbaranes(View):

    def get(self, request):
        form = CargarOferta()
        msg = ''
        titulo = 'Generador de Pre-Albaranes'
        return render(request, 'index.html', {'form': form, 'mensaje': msg, 'titulo': titulo})

    def post(self, request):
        if request.POST:
            form = CargarOferta(request.POST, request.FILES)
            if form.is_valid():
                archivo_albaran = form.cleaned_data.get('oferta')

                with open('csvofertas/albaran.csv', 'wb+') as destination:
                    for chunk in archivo_albaran.chunks():
                        destination.write(chunk)

                albaran = ''
                fecha = ''
                plazo = ''
                cliente = ''
                agente = ''
                proveedor = ''
                rsoc = ''
                empresa = ''
                dir = ''
                cp = ''
                pob = ''
                pro = ''
                tel = ''
                mail = ''
                icoterm = ''
                peso = ''
                importe_bruto = ''
                portes = ''
                imp_portes = ''
                dtopp = ''
                imp_dtopp = ''
                base_imp = ''
                iva = ''
                imp_iva = ''
                rec_quiv = ''
                imp_rec_quiv = ''
                total = ''
                forma_pago = ''
                giros = ''
                dp1 = ''
                dp2 = ''
                dp3 = ''
                transportista = ''
                iban = 'ES25 2100-1083-1102-0005-4013'

                doc = DocxTemplate("csvofertas/plantilla_pedido.docx")

                with open('csvofertas/albaran.csv') as csv_file:
                    csv_reader = csv.reader(csv_file, delimiter=';')
                    line_count = 0

                    for row in csv_reader:
                        if line_count == 1:
                            albaran = row[0]
                            fecha = row[1]
                            cliente = row[2]
                            proveedor = row[4]
                            rsoc = row[5]
                            empresa = row[6]
                            dir = row[7]
                            cp = row[8]
                            pob = row[9]
                            pro = row[10]
                            tel = row[11]
                            mail = row[13]
                            nif = row[15]
                            icoterm = row[17]
                            agente = '(' + row[18] + ')'
                            forma_pago = row[24]
                            giros = row[25]
                            dp1 = row[28]
                            dp2 = row[29]
                            dp3 = row[30]
                            transportista = row[35]
                            peso = row[48] + ' kg.'
                            importe_bruto = row[51]
                            portes = row[52]
                            imp_portes = row[53]
                            dtopp = row[55]
                            imp_dtopp = row[56]
                            base_imp = row[59]
                            iva = row[60]
                            imp_iva = row[61]
                            rec_quiv = row[62]
                            imp_rec_quiv = row[63]
                            total = row[64]
                        line_count += 1

                context = \
                    {
                        'PEDIDO': albaran,
                        'FECHA': fecha,
                        'CLIENTE': cliente,
                        'PROVEEDOR': proveedor,
                        'AGENTE': agente,
                        'PLAZOENTREGA': plazo,
                        'RSOC': rsoc,
                        'EMPRESA': empresa,
                        'DIR': dir,
                        'CP': cp,
                        'POB': pob,
                        'PRO': pro,
                        'NIF': nif,
                        'TEL': tel,
                        'MAIL': mail,
                    }

                doc.render(context)
                nombre_albaran = str(archivo_albaran).split('.')[0]
                ruta_guardado = 'C:/generador/pre-albaranes/' + nombre_albaran + '.docx'
                doc.save(ruta_guardado)

                doc = docx.Document(ruta_guardado)

                table = doc.add_table(rows=1, cols=6)

                for i in range(6):
                    for cell in table.columns[i].cells:
                        if i == 0:
                            cell.width = Cm(2.68)
                        elif i == 1:
                            cell.width = Cm(19)
                        elif i == 2:
                            cell.width = Inches(0.2)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        elif i == 3:
                            cell.width = Inches(0.2)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        elif i == 4:
                            cell.width = Inches(0.03)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        elif i == 5:
                            cell.width = Inches(0.2)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr = table.rows[0]
                hdr_cells = hdr.cells

                hdr_cells[0].paragraphs[0].add_run('REF.\n').font.size = Pt(9)
                hdr_cells[0].paragraphs[0].add_run('REF.\n').font.size = Pt(9)
                hdr_cells[0].paragraphs[0].runs[0].font.bold = True
                hdr_cells[0].paragraphs[0].runs[1].font.italic = True

                hdr_cells[1].paragraphs[0].add_run('DESCRIPCION\n').font.size = Pt(9)
                hdr_cells[1].paragraphs[0].add_run('SPECIFICATION\n').font.size = Pt(9)
                hdr_cells[1].paragraphs[0].runs[0].font.bold = True
                hdr_cells[1].paragraphs[0].runs[1].font.italic = True
                hdr_cells[1].paragraphs[0].runs[1].font.bold = False

                hdr_cells[2].paragraphs[0].add_run('CANTIDAD\n').font.size = Pt(9)
                hdr_cells[2].paragraphs[0].add_run('QUANTITY\n').font.size = Pt(9)
                hdr_cells[2].paragraphs[0].runs[0].font.bold = True
                hdr_cells[2].paragraphs[0].runs[1].font.italic = True
                hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr_cells[3].paragraphs[0].add_run('PRECIO\n').font.size = Pt(9)
                hdr_cells[3].paragraphs[0].add_run('PRICE\n').font.size = Pt(9)
                hdr_cells[3].paragraphs[0].add_run('EUROx100').font.size = Pt(9)
                hdr_cells[3].paragraphs[0].runs[0].font.bold = True
                hdr_cells[3].paragraphs[0].runs[1].font.italic = True
                hdr_cells[3].paragraphs[0].runs[2].font.bold = True
                hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr_cells[4].paragraphs[0].add_run('DTO.\n').font.size = Pt(9)
                hdr_cells[4].paragraphs[0].add_run('DIS.\n').font.size = Pt(9)
                hdr_cells[4].paragraphs[0].add_run('%').font.size = Pt(9)
                hdr_cells[4].paragraphs[0].runs[0].font.bold = True
                hdr_cells[4].paragraphs[0].runs[1].font.italic = True
                hdr_cells[4].paragraphs[0].runs[2].font.bold = True
                hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr_cells[5].paragraphs[0].add_run('IMPORTE\n').font.size = Pt(9)
                hdr_cells[5].paragraphs[0].add_run('AMOUNT\n').font.size = Pt(9)
                hdr_cells[5].paragraphs[0].add_run('EURO').font.size = Pt(9)
                hdr_cells[5].paragraphs[0].runs[0].font.bold = True
                hdr_cells[5].paragraphs[0].runs[1].font.italic = True
                hdr_cells[5].paragraphs[0].runs[2].font.bold = True
                hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr.height = Cm(1.25)
                hdr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                barra_cabeza = table.add_row()
                barra_cabeza_tabla = barra_cabeza.cells

                barra_cabeza_tabla[5].merge(barra_cabeza_tabla[4])
                barra_cabeza_tabla[4].merge(barra_cabeza_tabla[3])
                barra_cabeza_tabla[3].merge(barra_cabeza_tabla[2])
                barra_cabeza_tabla[2].merge(barra_cabeza_tabla[1])
                barra_cabeza_tabla[1].merge(barra_cabeza_tabla[0])

                barra_cabeza.height = Cm(0.65)
                barra_cabeza.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                insertHR(barra_cabeza_tabla[0].paragraphs[0])

                set_repeat_table_header(table.rows[0])
                set_repeat_table_header(table.rows[1])

                with open('csvofertas/albaran.csv') as csv_file:
                    csv_reader = csv.reader(csv_file, delimiter=';')
                    count = 0

                    for row in csv_reader:
                        if count > 2:
                            row_prod = table.add_row()
                            row_cells = row_prod.cells

                            row_prod.height = Cm(1)
                            row_prod.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                            row_cells[0].paragraphs[0].add_run(row[24]).font.size = Pt(9)
                            row_cells[0].paragraphs[0].add_run('\n' + row[3]).font.size = Pt(9)
                            row_cells[0].paragraphs[0].runs[1].font.italic = True
                            row_cells[0].paragraphs[0].runs[1].font.bold = True

                            row_cells[1].paragraphs[0].add_run(row[25]).font.size = Pt(8.5)
                            row_cells[1].paragraphs[0].add_run('\nPedido: ').font.size = Pt(8)
                            row_cells[1].paragraphs[0].add_run(row[2]).font.size = Pt(8)
                            row_cells[1].paragraphs[0].runs[2].font.bold = True

                            row_cells[2].text = row[12]
                            row_cells[2].paragraphs[0].runs[0].font.size = Pt(10)
                            row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            row_cells[3].text = row[13]
                            row_cells[3].paragraphs[0].runs[0].font.size = Pt(10)
                            row_cells[3].paragraphs[0].runs[0].font.bold = True
                            row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            if str(row[15]).strip() == '0' or str(row[15]).strip() == '':
                                row_cells[4].text = 'Neto'
                                row_cells[4].paragraphs[0].runs[0].font.size = Pt(10)
                                row_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                            else:
                                row_cells[4].text = row[15]
                                row_cells[4].paragraphs[0].runs[0].font.size = Pt(10)
                                row_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            row_cells[5].text = row[16]
                            row_cells[5].paragraphs[0].runs[0].font.size = Pt(10)
                            row_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                        count += 1

                barra_pie = table.add_row()
                barra_pie_tabla = barra_pie.cells

                barra_pie_tabla[5].merge(barra_pie_tabla[4])
                barra_pie_tabla[4].merge(barra_pie_tabla[3])
                barra_pie_tabla[3].merge(barra_pie_tabla[2])
                barra_pie_tabla[2].merge(barra_pie_tabla[1])
                barra_pie_tabla[1].merge(barra_pie_tabla[0])

                barra_pie.height = Cm(0.65)
                barra_pie.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                insertHR(barra_pie_tabla[0].paragraphs[0])

                doc.add_paragraph()
                doc.add_paragraph()

                table_resumen = doc.add_table(rows=13, cols=6)

                for i in range(5):
                    for cell in table_resumen.columns[i].cells:
                        if i == 0:
                            cell.width = Cm(2.49)
                        if i == 1:
                            cell.width = Cm(4.5)
                        if i == 2:
                            cell.width = Cm(1)
                        if i == 3:
                            cell.width = Cm(5.25)
                        if i == 4:
                            cell.width = Cm(1.73)
                        if i == 5:
                            cell.width = Cm(3.08)

                i = 0

                for row in table_resumen.rows:
                    if i == 7:
                        row.height = Cm(0.74)
                        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                    else:
                        row.height = Cm(0.35)
                        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                    i += 1

                '''Información del pedido'''
                if str(icoterm).strip() != '':
                    table_resumen.cell(0, 0).paragraphs[0].add_run('ICOTERM.').font.size = Pt(8)
                    table_resumen.cell(0, 0).paragraphs[0].runs[0].font.bold = True
                    table_resumen.cell(0, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(2, 0).paragraphs[0].add_run('PORTES').font.size = Pt(8)
                table_resumen.cell(2, 0).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(2, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(3, 0).paragraphs[0].add_run('TRANSPORT').font.size = Pt(8)
                table_resumen.cell(3, 0).paragraphs[0].runs[0].font.italic = True
                table_resumen.cell(3, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(5, 0).paragraphs[0].add_run('TRANSPORTE').font.size = Pt(8)
                table_resumen.cell(5, 0).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(5, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(6, 0).paragraphs[0].add_run('TRANSPORT').font.size = Pt(8)
                table_resumen.cell(6, 0).paragraphs[0].runs[0].font.italic = True
                table_resumen.cell(6, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(8, 0).paragraphs[0].add_run('PESO').font.size = Pt(8)
                table_resumen.cell(8, 0).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(8, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(9, 0).paragraphs[0].add_run('WEIGHT').font.size = Pt(8)
                table_resumen.cell(9, 0).paragraphs[0].runs[0].font.italic = True
                table_resumen.cell(9, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(0, 1).paragraphs[0].add_run(icoterm).font.size = Pt(8)

                if str(portes).strip() == 'D':
                    table_resumen.cell(2, 1).paragraphs[0].add_run('Portes debidos')
                    table_resumen.cell(3, 1).paragraphs[0].add_run('Transport not included')
                elif str(portes).strip() == 'P':
                    table_resumen.cell(2, 1).paragraphs[0].add_run('Portes pagados')
                    table_resumen.cell(3, 1).paragraphs[0].add_run('Transport included')
                elif str(portes).strip() == 'F':
                    table_resumen.cell(2, 1).paragraphs[0].add_run('Portes en factura')
                    table_resumen.cell(3, 1).paragraphs[0].add_run('Transport in invoice')
                else:
                    table_resumen.cell(2, 1).paragraphs[0].text = 'A concretar'
                    table_resumen.cell(3, 1).paragraphs[0].text = 'To be determined'

                table_resumen.cell(2, 1).paragraphs[0].runs[0].font.size = Pt(8)
                table_resumen.cell(3, 1).paragraphs[0].runs[0].font.size = Pt(8)
                table_resumen.cell(2, 1).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(3, 1).paragraphs[0].runs[0].font.italic = True

                if transportista.strip() == '':
                    table_resumen.cell(5, 1).paragraphs[0].add_run('A concretar\n').font.size = Pt(8)
                    table_resumen.cell(5, 1).paragraphs[0].runs[0].font.bold = True
                    table_resumen.cell(6, 1).paragraphs[0].add_run('To be determined').font.size = Pt(8)
                    table_resumen.cell(6, 1).paragraphs[0].runs[0].font.italic = True
                else:
                    table_resumen.cell(5, 1).paragraphs[0].text = transportista
                    table_resumen.cell(5, 1).paragraphs[0].runs[0].font.size = Pt(8)

                table_resumen.cell(8, 1).paragraphs[0].text = peso
                table_resumen.cell(8, 1).paragraphs[0].runs[0].font.size = Pt(8)

                # Resumen del pedido

                table_resumen.cell(1, 3).paragraphs[0].add_run('IMPORTE BRUTO / ').font.size = Pt(8)
                table_resumen.cell(1, 3).paragraphs[0].add_run('GROSS AMOUNT').font.size = Pt(8)
                table_resumen.cell(1, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(1, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(1, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(1, 5).paragraphs[0].add_run(importe_bruto + ' €').font.size = Pt(8)
                table_resumen.cell(1, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(2, 3).paragraphs[0].add_run('PORTES / ').font.size = Pt(8)
                table_resumen.cell(2, 3).paragraphs[0].add_run('TRANSPORT').font.size = Pt(8)
                table_resumen.cell(2, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(2, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(2, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(2, 5).paragraphs[0].add_run(imp_portes + ' €').font.size = Pt(8)
                table_resumen.cell(2, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(3, 3).paragraphs[0].add_run('DTOP. PP.').font.size = Pt(8)
                table_resumen.cell(3, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(3, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(3, 4).paragraphs[0].add_run(dtopp + ' %').font.size = Pt(8)
                table_resumen.cell(3, 4).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(3, 5).paragraphs[0].add_run(imp_dtopp + ' €').font.size = Pt(8)
                table_resumen.cell(3, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(4, 3).paragraphs[0].add_run('BASE IMPONIBLE / ').font.size = Pt(8)
                table_resumen.cell(4, 3).paragraphs[0].add_run('TAXABLE BASE').font.size = Pt(8)
                table_resumen.cell(4, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(4, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(4, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(4, 5).paragraphs[0].add_run(base_imp + ' €').font.size = Pt(8)
                table_resumen.cell(4, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(5, 3).paragraphs[0].add_run('IVA / ').font.size = Pt(8)
                table_resumen.cell(5, 3).paragraphs[0].add_run('IVA').font.size = Pt(8)
                table_resumen.cell(5, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(5, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(5, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(5, 4).paragraphs[0].add_run(iva + ' %').font.size = Pt(8)
                table_resumen.cell(5, 4).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(5, 5).paragraphs[0].add_run(imp_iva + ' €').font.size = Pt(8)
                table_resumen.cell(5, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(6, 3).paragraphs[0].add_run('REC. EQUIVALENCIA').font.size = Pt(8)
                table_resumen.cell(6, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(6, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(6, 4).paragraphs[0].add_run(rec_quiv + ' %').font.size = Pt(8)
                table_resumen.cell(6, 4).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(6, 5).paragraphs[0].add_run(imp_rec_quiv + ' €').font.size = Pt(8)
                table_resumen.cell(6, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(7, 3).merge(table_resumen.cell(7, 4))
                table_resumen.cell(7, 4).merge(table_resumen.cell(7, 5))
                insertHR(table_resumen.cell(7, 3).paragraphs[0])

                table_resumen.cell(8, 3).paragraphs[0].add_run('IMPORTE TOTAL / ').font.size = Pt(9)
                table_resumen.cell(8, 3).paragraphs[0].add_run('TOTAL AMOUNT').font.size = Pt(9)
                table_resumen.cell(8, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(8, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(8, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(8, 5).paragraphs[0].add_run(total + ' €').font.size = Pt(9)
                table_resumen.cell(8, 5).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(8, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                table_resumen.cell(11, 3).paragraphs[0].add_run('FORMA DE PAGO / ').font.size = Pt(8)
                table_resumen.cell(11, 3).paragraphs[0].add_run('MEANS OF PAYMENT').font.size = Pt(8)
                table_resumen.cell(11, 3).paragraphs[0].runs[0].font.bold = True
                table_resumen.cell(11, 3).paragraphs[0].runs[1].font.italic = True
                table_resumen.cell(11, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                table_resumen.cell(12, 5).merge(table_resumen.cell(12, 4))

                if str(forma_pago).strip() == 'TRANSFERENCIA' or str(forma_pago).strip() == 'CONTADO':
                    table_resumen.cell(11, 5).paragraphs[0].add_run(str(forma_pago).strip()).font.size = Pt(8)
                    table_resumen.cell(11, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    table_resumen.cell(12, 3).paragraphs[0].add_run('SWIFT/IBAN').font.size = Pt(9)
                    table_resumen.cell(12, 3).paragraphs[0].runs[0].font.bold = True
                    table_resumen.cell(12, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                    table_resumen.cell(12, 5).paragraphs[0].add_run(iban).font.size = Pt(9)
                    table_resumen.cell(12, 5).paragraphs[0].runs[0].font.bold = True
                    table_resumen.cell(12, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                elif str(forma_pago).strip() == 'GIRO':
                    table_resumen.cell(11, 5).paragraphs[0].add_run(
                        str(forma_pago).strip() + ' a ' + giros + ' DIAS').font.size = Pt(8)
                    table_resumen.cell(11, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    table_resumen.cell(12, 5).paragraphs[0].add_run('DIAS ').font.size = Pt(8)

                    if str(dp1).strip() != '0' or str(dp1).strip() != '':
                        table_resumen.cell(12, 5).paragraphs[0].add_run(dp1).font.size = Pt(8)
                        table_resumen.cell(12, 5).paragraphs[0].runs[1].font.bold = True
                        table_resumen.cell(12, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    if str(dp2).strip() != '0' or str(dp1).strip() != '':
                        table_resumen.cell(12, 5).paragraphs[0].add_run('/' + dp2).font.size = Pt(8)
                        table_resumen.cell(12, 5).paragraphs[0].runs[2].font.bold = True
                        table_resumen.cell(12, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    if str(dp3).strip() != '0' or str(dp1).strip() != '':
                        table_resumen.cell(12, 5).paragraphs[0].add_run('/' + dp3).font.size = Pt(8)
                        table_resumen.cell(12, 5).paragraphs[0].runs[3].font.bold = True
                        table_resumen.cell(12, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                else:
                    table_resumen.cell(11, 5).paragraphs[0].add_run(str(forma_pago).strip()).font.size = Pt(8)
                    table_resumen.cell(11, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                doc.save(ruta_guardado)

                # os.startfile(ruta_guardado)

                return redirect('pre-albaranes')
            else:
                form = CargarOferta()
                msg = 'Fichero no válido. Porfavor, compruebe el archivo.'
                titulo = 'Generador de Pre-Albaranes'
                return render(request, 'index.html', {'form': form, 'mensaje': msg, 'titulo': titulo})


class Consultas(View):
    def get(self, request):
        form = CargarOferta()
        msg = ''
        titulo = 'Consultas a proveedores'
        return render(request, 'index.html', {'form': form, 'mensaje': msg, 'titulo': titulo})

    def post(self, request):
        if request.POST:
            form = CargarOferta(request.POST, request.FILES)
            if form.is_valid():
                archivo_consulta = form.cleaned_data.get('oferta')

                with open('csvofertas/consulta.csv', 'wb+') as destination:
                    for chunk in archivo_consulta.chunks():
                        destination.write(chunk)

                consulta = ''
                fecha = ''
                validez = ''
                proveedor = ''
                rsoc = ''
                empresa = ''
                dir = ''
                cp = ''
                pob = ''
                pro = ''
                tel = ''
                observaciones = ''

                doc = DocxTemplate("csvofertas/plantilla_cons.docx")

                with open('csvofertas/consulta.csv') as csv_file:
                    csv_reader = csv.reader(csv_file, delimiter=';')
                    line_count = 0

                    for row in csv_reader:
                        if line_count == 1:
                            consulta = row[0]
                            fecha = row[1]
                            validez = row[2]
                            proveedor = row[3]
                            rsoc = row[5]
                            empresa = row[6]
                            dir = row[7]
                            cp = row[8]
                            pob = row[9]
                            pro = row[10]
                            tel = row[11]
                            observaciones = str(row[13]) + str(row[14]) + str(row[15]) + str(row[16])
                            break
                        line_count += 1

                context = \
                    {
                        'CONSULTA': consulta,
                        'FECHA': fecha,
                        'VALIDEZ': validez,
                        'PROVEEDOR': proveedor,
                        'RSOC': rsoc,
                        'EMPRESA': empresa,
                        'DIR': dir,
                        'CP': cp,
                        'POB': pob,
                        'PRO': pro,
                        'TEL': tel,
                        # 'MAIL': mail,
                    }

                doc.render(context)

                nombre_consulta = str(archivo_consulta).split('.')[0]
                ruta_guardado = 'C:/generador/consultas/' + nombre_consulta + '.docx'
                doc.save(ruta_guardado)

                doc = docx.Document(ruta_guardado)

                condiciones = doc.add_paragraph()
                condiciones.add_run(
                    'Muy Sres. Nuestros: \nRogamos nos envíen su mejor precio y plazo de entrega para los siguientes artículos.\nPlease, send us your best price and delivery date for the following references.').font.size = Pt(
                    11)

                table = doc.add_table(rows=1, cols=4)

                for i in range(4):
                    for cell in table.columns[i].cells:
                        if i == 0:
                            cell.width = Inches(4)
                        elif i == 1:
                            cell.width = Inches(1)
                        elif i == 2:
                            cell.width = Inches(0.5)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        elif i == 3:
                            cell.width = Inches(0.5)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr = table.rows[0]
                hdr_cells = hdr.cells

                hdr_cells[0].paragraphs[0].add_run('DESCRIPCION\n').font.size = Pt(9)
                hdr_cells[0].paragraphs[0].add_run('SPECIFICATION\n').font.size = Pt(9)
                hdr_cells[0].paragraphs[0].runs[0].font.bold = True
                hdr_cells[0].paragraphs[0].runs[1].font.italic = True
                hdr_cells[0].paragraphs[0].runs[1].font.bold = False

                hdr_cells[1].paragraphs[0].add_run('CANTIDAD\n').font.size = Pt(9)
                hdr_cells[1].paragraphs[0].add_run('QUANTITY\n').font.size = Pt(9)
                hdr_cells[1].paragraphs[0].runs[0].font.bold = True
                hdr_cells[1].paragraphs[0].runs[1].font.italic = True
                hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr_cells[2].paragraphs[0].add_run('PRECIO\n').font.size = Pt(9)
                hdr_cells[2].paragraphs[0].add_run('PRICE\n').font.size = Pt(9)
                hdr_cells[2].paragraphs[0].add_run('EUROx100').font.size = Pt(9)
                hdr_cells[2].paragraphs[0].runs[0].font.bold = True
                hdr_cells[2].paragraphs[0].runs[1].font.italic = True
                hdr_cells[2].paragraphs[0].runs[2].font.bold = True
                hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr_cells[3].paragraphs[0].add_run('DTO.\n').font.size = Pt(9)
                hdr_cells[3].paragraphs[0].add_run('DIS.\n').font.size = Pt(9)
                hdr_cells[3].paragraphs[0].add_run('%').font.size = Pt(9)
                hdr_cells[3].paragraphs[0].runs[0].font.bold = True
                hdr_cells[3].paragraphs[0].runs[1].font.italic = True
                hdr_cells[3].paragraphs[0].runs[2].font.bold = True
                hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr.height = Cm(1.25)
                hdr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                barra_cabeza = table.add_row()
                barra_cabeza_tabla = barra_cabeza.cells

                barra_cabeza_tabla[3].merge(barra_cabeza_tabla[2])
                barra_cabeza_tabla[2].merge(barra_cabeza_tabla[1])
                barra_cabeza_tabla[1].merge(barra_cabeza_tabla[0])

                barra_cabeza.height = Cm(0.65)
                barra_cabeza.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                insertHR(barra_cabeza_tabla[0].paragraphs[0])

                set_repeat_table_header(table.rows[0])
                set_repeat_table_header(table.rows[1])

                with open('csvofertas/consulta.csv') as csv_file:
                    csv_reader = csv.reader(csv_file, delimiter=';')
                    count = 0

                    for row in csv_reader:
                        if count > 2:
                            row_prod = table.add_row()
                            row_cells = row_prod.cells

                            row_prod.height = Cm(1)
                            row_prod.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                            row_cells[0].paragraphs[0].add_run(
                                str(row[8]).strip() + str(row[9]).strip() + str(row[10]).strip() + str(
                                    row[11]).strip()).font.size = Pt(10)
                            row_cells[0].paragraphs[0].add_run('\nRef. ' + row[7]).font.size = Pt(10)
                            row_cells[0].paragraphs[0].runs[1].font.italic = True

                            row_cells[1].text = row[12]
                            row_cells[1].paragraphs[0].runs[0].font.size = Pt(10)
                            row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                        count += 1

                barra_pie = table.add_row()
                barra_pie_tabla = barra_pie.cells

                barra_pie_tabla[3].merge(barra_pie_tabla[2])
                barra_pie_tabla[2].merge(barra_pie_tabla[1])
                barra_pie_tabla[1].merge(barra_pie_tabla[0])

                barra_pie.height = Cm(0.65)
                barra_pie.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                insertHR(barra_pie_tabla[0].paragraphs[0])

                condiciones = doc.add_paragraph()
                condiciones.add_run('\n\n\tROGAMOS NOS COMINIQUEN EL PAIS DE ORIGEN Y TARIC').font.size = Pt(11)

                obs = doc.add_paragraph()
                obs.add_run('\n' + observaciones).font.size = Pt(11)

                doc.save(ruta_guardado)

                # os.startfile(ruta_guardado)

                return redirect('consultas')
            else:
                form = CargarOferta()
                msg = 'Fichero no válido. Porfavor, compruebe el archivo.'
                titulo = 'Consultas a proveedores'
                return render(request, 'index.html', {'form': form, 'mensaje': msg, 'titulo': titulo})


class PedidosProv(View):
    def get(self, request):
        form = CargarOferta()
        msg = ''
        titulo = 'Pedidos a proveedores'
        return render(request, 'index.html', {'form': form, 'mensaje': msg, 'titulo': titulo})

    def post(self, request):
        if request.POST:
            form = CargarOferta(request.POST, request.FILES)
            if form.is_valid():
                archivo_pedido = form.cleaned_data.get('oferta')

                with open('csvofertas/pedprov.csv', 'wb+') as destination:
                    for chunk in archivo_pedido.chunks():
                        destination.write(chunk)

                pedido = ''
                fecha = ''
                entrega = ''
                proveedor = ''
                rsoc = ''
                empresa = ''
                dir = ''
                cp = ''
                pob = ''
                pro = ''
                tel = ''
                mail = ''
                peso = ''
                observaciones = ''

                doc = DocxTemplate("csvofertas/plantilla_pedprov.docx")

                with open('csvofertas/pedprov.csv') as csv_file:
                    csv_reader = csv.reader(csv_file, delimiter=';')
                    line_count = 0

                    for row in csv_reader:
                        if line_count == 1:
                            pedido = row[0]
                            fecha = row[1]
                            entrega = row[2]
                            proveedor = row[3]
                            rsoc = row[5]
                            empresa = row[6]
                            dir = row[7]
                            cp = row[8]
                            pob = row[9]
                            pro = row[10]
                            tel = row[11]
                            mail = row[13]
                            peso = str(row[18]).replace(',', '.')
                            observaciones = str(row[21]) + str(row[22]) + str(row[23]) + str(row[24])
                            break
                        line_count += 1

                context = \
                    {
                        'PEDIDO': pedido,
                        'FECHA': fecha,
                        'ENTREGA': entrega,
                        'PROVEEDOR': proveedor,
                        'RSOC': rsoc,
                        'EMPRESA': empresa,
                        'DIR': dir,
                        'CP': cp,
                        'POB': pob,
                        'PRO': pro,
                        'TEL': tel,
                        'MAIL': mail,
                    }

                doc.render(context)

                nombre_pedido = str(archivo_pedido).split('.')[0]
                ruta_guardado = 'C:/generador/ped_proveedores/' + nombre_pedido + '.docx'
                doc.save(ruta_guardado)

                doc = docx.Document(ruta_guardado)

                doc.add_paragraph()

                table = doc.add_table(rows=1, cols=5)

                for i in range(5):
                    for cell in table.columns[i].cells:
                        if i == 0:
                            cell.width = Inches(4)
                        elif i == 1:
                            cell.width = Inches(1)
                        elif i == 2:
                            cell.width = Inches(0.5)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        elif i == 3:
                            cell.width = Inches(0.5)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        elif i == 4:
                            cell.width = Inches(1.5)
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr = table.rows[0]
                hdr_cells = hdr.cells

                hdr_cells[0].paragraphs[0].add_run('DESCRIPCION\n').font.size = Pt(9)
                hdr_cells[0].paragraphs[0].add_run('SPECIFICATION\n').font.size = Pt(9)
                hdr_cells[0].paragraphs[0].runs[0].font.bold = True
                hdr_cells[0].paragraphs[0].runs[1].font.italic = True
                hdr_cells[0].paragraphs[0].runs[1].font.bold = False

                hdr_cells[1].paragraphs[0].add_run('CANTIDAD\n').font.size = Pt(9)
                hdr_cells[1].paragraphs[0].add_run('QUANTITY\n').font.size = Pt(9)
                hdr_cells[1].paragraphs[0].runs[0].font.bold = True
                hdr_cells[1].paragraphs[0].runs[1].font.italic = True
                hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr_cells[2].paragraphs[0].add_run('PRECIO\n').font.size = Pt(9)
                hdr_cells[2].paragraphs[0].add_run('PRICE\n').font.size = Pt(9)
                hdr_cells[2].paragraphs[0].add_run('EUROx100').font.size = Pt(9)
                hdr_cells[2].paragraphs[0].runs[0].font.bold = True
                hdr_cells[2].paragraphs[0].runs[1].font.italic = True
                hdr_cells[2].paragraphs[0].runs[2].font.bold = True
                hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr_cells[3].paragraphs[0].add_run('DTO.\n').font.size = Pt(9)
                hdr_cells[3].paragraphs[0].add_run('DIS.\n').font.size = Pt(9)
                hdr_cells[3].paragraphs[0].add_run('%').font.size = Pt(9)
                hdr_cells[3].paragraphs[0].runs[0].font.bold = True
                hdr_cells[3].paragraphs[0].runs[1].font.italic = True
                hdr_cells[3].paragraphs[0].runs[2].font.bold = True
                hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr_cells[4].paragraphs[0].add_run('IMPORTE\n').font.size = Pt(9)
                hdr_cells[4].paragraphs[0].add_run('AMOUNT\n').font.size = Pt(9)
                hdr_cells[4].paragraphs[0].add_run('EURO').font.size = Pt(9)
                hdr_cells[4].paragraphs[0].runs[0].font.bold = True
                hdr_cells[4].paragraphs[0].runs[1].font.italic = True
                hdr_cells[4].paragraphs[0].runs[2].font.bold = True
                hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                hdr.height = Cm(1.25)
                hdr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                barra_cabeza = table.add_row()
                barra_cabeza_tabla = barra_cabeza.cells

                barra_cabeza_tabla[4].merge(barra_cabeza_tabla[3])
                barra_cabeza_tabla[3].merge(barra_cabeza_tabla[2])
                barra_cabeza_tabla[2].merge(barra_cabeza_tabla[1])
                barra_cabeza_tabla[1].merge(barra_cabeza_tabla[0])

                barra_cabeza.height = Cm(0.65)
                barra_cabeza.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                insertHR(barra_cabeza_tabla[0].paragraphs[0])

                set_repeat_table_header(table.rows[0])
                set_repeat_table_header(table.rows[1])

                with open('csvofertas/pedprov.csv') as csv_file:
                    csv_reader = csv.reader(csv_file, delimiter=';')
                    count = 0

                    for row in csv_reader:
                        if count > 2:
                            if str(row[9]).strip() != 'Texto':
                                row_prod = table.add_row()
                                row_cells = row_prod.cells

                                row_prod.height = Cm(1)
                                row_prod.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                                row_cells[0].paragraphs[0].add_run(
                                    str(row[8]).strip() + str(row[9]).strip() + str(row[10]).strip() + str(
                                        row[11]).strip()).font.size = Pt(10)
                                row_cells[0].paragraphs[0].add_run('\nRef. ' + row[7]).font.size = Pt(10)
                                row_cells[0].paragraphs[0].runs[1].font.italic = True

                                row_cells[1].text = row[12]
                                row_cells[1].paragraphs[0].runs[0].font.size = Pt(10)
                                row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                                row_cells[2].text = row[14]
                                row_cells[2].paragraphs[0].runs[0].font.size = Pt(10)
                                row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                                if int(float(str(row[15]).replace(',', '.'))) == 0:
                                    row_cells[3].text = 'Neto'
                                else:
                                    row_cells[3].text = row[15]
                                row_cells[3].paragraphs[0].runs[0].font.size = Pt(10)
                                row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                                row_cells[4].text = row[16]
                                row_cells[4].paragraphs[0].runs[0].font.size = Pt(10)
                                row_cells[4].paragraphs[0].runs[0].font.bold = True
                                row_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                            else:
                                row_prod = table.add_row()
                                row_cells = row_prod.cells

                                row_cells[1].merge(row_cells[0])
                                row_cells[0].paragraphs[0].add_run(row[8]).font.size = Pt(10)
                                row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        count += 1

                row_peso = table.add_row()
                row_cells = row_peso.cells

                row_cells[1].merge(row_cells[0])

                row_cells[0].paragraphs[0].add_run('PESO/Weight:\t' + str(peso) + ' Kg.').font.size = Pt(10)
                row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                barra_pie = table.add_row()
                barra_pie_tabla = barra_pie.cells

                barra_pie_tabla[4].merge(barra_pie_tabla[3])
                barra_pie_tabla[3].merge(barra_pie_tabla[2])
                barra_pie_tabla[2].merge(barra_pie_tabla[1])
                barra_pie_tabla[1].merge(barra_pie_tabla[0])

                barra_pie.height = Cm(0.65)
                barra_pie.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                insertHR(barra_pie_tabla[0].paragraphs[0])

                obs = doc.add_paragraph()
                obs.add_run('\n' + observaciones).font.size = Pt(11)

                doc.save(ruta_guardado)

                # os.startfile(ruta_guardado)

                return redirect('pedidos-prov')

        else:
            form = CargarOferta()
            msg = 'Fichero no válido. Porfavor, compruebe el archivo.'
            titulo = 'Pedidos a proveedores'
            return render(request, 'index.html', {'form': form, 'mensaje': msg, 'titulo': titulo})
